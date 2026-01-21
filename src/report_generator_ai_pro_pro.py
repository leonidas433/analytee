# autos: Leox433
# Metodo: AI
# Version: V2.1
# Fecha de última modificación: 2025-10-19 23:05 UTC

import os, re, shutil, unicodedata, time
from pathlib import Path
from datetime import datetime
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

PRIMARY = RGBColor(0x00, 0x4E, 0x89)
TEXT = RGBColor(0x2C, 0x2C, 0x2C)
BG_SOFT = "E7F0FA"

def _normalize_client_name(name: str) -> str:
    name = re.sub(r'_ChIJ[\w-]+$', '', name)
    name = name.replace('_', ' ').strip()
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("utf-8")
    name = re.sub(r'[^A-Za-z0-9 áéíóúÁÉÍÓÚüÜñÑ-]', '', name)
    return name.strip()

def _insert_cover(doc, client_display, title):
    p = doc.add_paragraph(title)
    p.style = doc.styles['Title']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(f"Cliente: {client_display}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def _insert_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run("―" * 40)
    run.font.color.rgb = RGBColor(150, 150, 150)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _style_table(table):
    table.style = 'Table Grid'
    for cell in table.rows[0].cells:
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _rating_chart_by_star(df, path):
    try:
        counts = df['rating_num'].value_counts().sort_index()
        plt.figure(figsize=(6, 4))
        plt.bar(counts.index.astype(str), counts.values, color="#004E89", alpha=0.85, edgecolor="white")
        plt.xlabel("Estrellas")
        plt.ylabel("Cantidad")
        plt.title("Distribución de valoraciones")
        plt.tight_layout()
        plt.savefig(path, dpi=160)
        plt.close()
        return path
    except Exception as e:
        logger.error(f"Error creando gráfico rating: {e}")
        return None

def _sentiment_chart(df, path):
    try:
        order = ["muy_negativo", "negativo", "neutro", "positivo", "muy_positivo"]
        labels = ["Muy negativo", "Negativo", "Neutro", "Positivo", "Muy positivo"]
        colors = ["#C62828", "#E53935", "#FB8C00", "#43A047", "#2E7D32"]
        v = df['sentiment'].value_counts()
        vals = [int(v.get(k, 0)) for k in order]
        plt.figure(figsize=(6, 4))
        plt.barh(labels, vals, color=colors)
        plt.title("Distribución de sentimiento")
        plt.tight_layout()
        plt.savefig(path, dpi=160)
        plt.close()
        return path
    except Exception as e:
        logger.error(f"Error creando gráfico sentimiento: {e}")
        return None

def _kpi_card_row(doc, kpis):
    table = doc.add_table(rows=1, cols=len(kpis))
    table.autofit = False
    for i, (title, value, note) in enumerate(kpis):
        cell = table.cell(0, i)
        shade = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), BG_SOFT))
        cell._tc.get_or_add_tcPr().append(shade)

        p1 = cell.paragraphs[0]
        r1 = p1.add_run(title.upper())
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = PRIMARY
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = cell.add_paragraph(str(value))
        r2 = p2.runs[0]
        r2.font.bold = True
        r2.font.size = Pt(22)
        r2.font.color.rgb = TEXT
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p3 = cell.add_paragraph(note)
        for r in p3.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(100, 100, 100)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_report_ai_pro(df, client_name, output_path, cfg, project_root="."):
    try:
        client_clean = _normalize_client_name(client_name)
        safe_name = re.sub(r'[^A-Za-z0-9_-]', '', client_clean)
        client_dir = Path(cfg.get('output_dir', './reports')) / safe_name
        client_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()
        _insert_cover(doc, client_clean, cfg.get('report_title', 'Informe ORM y CX'))

        avg = df['rating_num'].mean() if len(df) else 0
        total = len(df)
        pos = (df['sentiment'].isin(['positivo', 'muy_positivo']).sum() / total * 100) if total else 0

        _kpi_card_row(doc, [
            ("Valoración media", f"{avg:.2f}", "Promedio global"),
            ("Total reseñas", f"{total}", "Volumen analizado"),
            ("% positivas", f"{pos:.0f}%", "Positivas + Muy positivas"),
        ])

        _insert_separator(doc)

        g1 = _rating_chart_by_star(df, client_dir / 'rating_chart.png')
        g2 = _sentiment_chart(df, client_dir / 'sentiment_chart.png')

        if g1 and Path(g1).exists():
            doc.add_picture(str(g1), width=Inches(5.5))
        if g2 and Path(g2).exists():
            doc.add_picture(str(g2), width=Inches(5.5))

        output_file = client_dir / f"{safe_name}_informe_IA_PRO.docx"
        doc.save(str(output_file))
        logger.info(f"✅ Informe guardado: {output_file}")
        return str(output_file)

    except Exception as e:
        logger.error(f"Error generando informe: {e}")
        raise
