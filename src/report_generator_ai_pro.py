# report_generator_ai_pro_fixed.py — Improved version with better error handling
# ORM Analyzer IA PRO — Fixed Version

import os, re, shutil, unicodedata, time
from pathlib import Path
from datetime import datetime
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

PRIMARY = RGBColor(0x00, 0x4E, 0x89)
TEXT = RGBColor(0x2C, 0x2C, 0x2C)
BG_SOFT = "F2F4F7"

# ========================== Utils ==========================
def _normalize_client_name(name: str) -> str:
    name = re.sub(r'_ChIJ[\w-]+$', '', name)
    name = name.replace('_', ' ').strip()
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("utf-8")
    name = re.sub(r'[^A-Za-z0-9 áéíóúÁÉÍÓÚüÜñÑ-]', '', name)
    return name.strip()

def _safe_filename(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9 _-]', '', name)

def _set_cell_shade(cell, hex_color):
    """Safely set cell background color"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Use simpler approach for shading
        from docx.oxml import parse_xml
        shade_xml = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>'
        )
        tcPr.append(shade_xml)
    except Exception as e:
        logger.warning(f"Could not set cell shading: {e}")

def _insert_simple_toc(doc, title="Contenido del informe"):
    """Insert a simple TOC without complex XML manipulation"""
    p_title = doc.add_paragraph(title)
    p_title.style = doc.styles['Heading 1']
    
    # Add a simple placeholder TOC that can be updated manually in Word
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("Tabla de contenido - Actualizar en Word: Clic derecho > Actualizar campo")
    toc_run.font.italic = True
    toc_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    return toc_para

# ========================== Doc styling ==========================
def _apply_doc_standards(doc, client_display):
    """Apply document formatting with better error handling"""
    try:
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

            # Header
            header = section.header.paragraphs[0]
            run = header.add_run(f"Informe ORM — {client_display}")
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = PRIMARY
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer - simplified without field codes
            footer = section.footer.paragraphs[0]
            footer_run = footer.add_run(f"{datetime.now().strftime('%Y-%m-%d')} • Página [Actualizar numeración en Word]")
            footer_run.font.name = "Arial"
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = TEXT

        # Normal style
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
        
    except Exception as e:
        logger.warning(f"Could not apply all document standards: {e}")

# ========================== Prompt helpers ==========================
def _load_prompt(project_root: Path) -> str:
    p = project_root / "src" / "prompt.txt"
    try:
        return p.read_text(encoding="utf-8-sig") if p.exists() else "{{reviews_text}}"
    except Exception as e:
        logger.warning(f"Could not load prompt: {e}")
        return "{{reviews_text}}"

def _build_reviews_text(df):
    rows = []
    for _, r in df.iterrows():
        try:
            user = str(r.get("user", "")).strip() or "Usuario"
            date = str(r.get("date", "")).strip() or "s/f"
            rating = int(r.get("rating_num", 0))
            txt = str(r.get("text", "")).strip() or "[Sin comentario]"
            rows.append(f"- [{date}] {user} ({rating}★): {txt}")
        except Exception as e:
            logger.warning(f"Error processing review row: {e}")
            continue
    return "\n".join(rows)

def _call_openai(prompt: str, model: str, temperature: float):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return "Informe generado sin conexión a IA (defina OPENAI_API_KEY)."
    
    try:
        client = OpenAI(api_key=api_key)
        r = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[{"role": "user", "content": prompt}],
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"OpenAI API error: {e}")
        return f"Error al generar contenido con IA: {e}"

# ========================== Charts & KPI ==========================
def _kpi_card_row(doc, kpis):
    try:
        table = doc.add_table(rows=1, cols=len(kpis))
        table.autofit = True
        
        for i, (title, value, note) in enumerate(kpis):
            cell = table.cell(0, i)
            _set_cell_shade(cell, BG_SOFT)
            
            p1 = cell.paragraphs[0]
            r1 = p1.add_run(title.upper())
            r1.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = PRIMARY
            
            p2 = cell.add_paragraph()
            r2 = p2.add_run(str(value))
            r2.bold = True
            r2.font.size = Pt(16)
            r2.font.color.rgb = TEXT
            
            if note:
                p3 = cell.add_paragraph(note)
                for r in p3.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = TEXT
    except Exception as e:
        logger.error(f"Error creating KPI table: {e}")

def _rating_chart_by_star(df, path):
    try:
        counts = df["rating_num"].value_counts().sort_index()
        plt.figure(figsize=(8, 6))
        plt.bar(counts.index.astype(str), counts.values, color="#4A90E2", edgecolor="#1F4E79")
        plt.xlabel("Estrellas")
        plt.ylabel("Cantidad")
        plt.title("Distribución por estrellas")
        plt.grid(axis="y", alpha=0.2)
        plt.tight_layout()
        plt.savefig(path, dpi=160, bbox_inches="tight")
        plt.close()
        return path
    except Exception as e:
        logger.error(f"Error creating rating chart: {e}")
        return None

def _sentiment_chart(df, path):
    try:
        order = ["muy_negativo", "negativo", "neutro", "positivo", "muy_positivo"]
        labels = ["Muy negativo", "Negativo", "Neutro", "Positivo", "Muy positivo"]
        colors = ["#C62828", "#E53935", "#FB8C00", "#43A047", "#2E7D32"]
        v = df["sentiment"].value_counts()
        vals = [int(v.get(k, 0)) for k in order]
        
        plt.figure(figsize=(8, 6))
        plt.barh(labels, vals, color=colors)
        plt.title("Distribución por sentimiento")
        plt.grid(axis="x", alpha=0.2)
        plt.tight_layout()
        plt.savefig(path, dpi=160, bbox_inches="tight")
        plt.close()
        return path
    except Exception as e:
        logger.error(f"Error creating sentiment chart: {e}")
        return None

# ========================== Save DOCX + PDF ==========================
def _save_word_compatible(doc, path):
    """Improved save function with better error handling"""
    path = Path(path)
    
    # Handle existing file
    if path.exists():
        backup = path.with_name(path.stem + f"_{datetime.now().strftime('%H%M%S')}.bak.docx")
        try:
            shutil.copy2(path, backup)
            logger.info(f"Backup created: {backup.name}")
        except Exception as e:
            logger.warning(f"Could not create backup: {e}")

    # Save DOCX
    try:
        doc.save(str(path))
        logger.info(f"✅ Documento Word guardado correctamente: {path}")
        time.sleep(2)  # Allow file system to settle
    except Exception as e:
        logger.error(f"Error saving DOCX: {e}")
        raise

    # Try multiple PDF conversion methods
    pdf_path = path.with_suffix(".pdf")
    pdf_success = False

    # Method 1: docx2pdf
    try:
        from docx2pdf import convert
        convert(str(path), str(pdf_path))
        logger.info(f"📄 PDF generado con docx2pdf: {pdf_path}")
        pdf_success = True
    except Exception as e:
        logger.warning(f"docx2pdf falló: {e}")

    # Method 2: python-docx2txt + reportlab (fallback)
    if not pdf_success:
        try:
            _convert_via_reportlab(path, pdf_path)
            logger.info(f"📄 PDF generado con reportlab: {pdf_path}")
            pdf_success = True
        except Exception as e:
            logger.warning(f"reportlab conversion falló: {e}")

    # Method 3: Microsoft Word COM (last resort)
    if not pdf_success:
        try:
            _convert_via_word_com(path, pdf_path)
            logger.info(f"📄 PDF generado con Microsoft Word: {pdf_path}")
            pdf_success = True
        except Exception as e:
            logger.warning(f"Word COM conversion falló: {e}")

    if not pdf_success:
        logger.warning("⚠️ No se pudo generar PDF - solo disponible DOCX")

    return str(path)

def _convert_via_reportlab(docx_path, pdf_path):
    """Alternative PDF conversion using reportlab"""
    try:
        import docx2txt
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        
        # Extract text from DOCX
        text = docx2txt.process(str(docx_path))
        
        # Create PDF
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
        
        doc.build(story)
        
    except ImportError:
        raise Exception("reportlab or docx2txt not available")

def _convert_via_word_com(docx_path, pdf_path):
    """Microsoft Word COM conversion with better error handling"""
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM
        pythoncom.CoInitialize()
        
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            # Open document
            doc = word.Documents.Open(str(docx_path.resolve()))
            
            # Save as PDF
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
            
            # Clean up
            doc.Close(False)
            word.Quit()
            
        finally:
            pythoncom.CoUninitialize()
            
    except ImportError:
        raise Exception("pywin32 not available")

# ========================== Main generator ==========================
def create_report_ai_pro(df, client_name, output_path, cfg, project_root="."):
    """Main report generation function with improved error handling"""
    try:
        project_root = Path(project_root)
        base_prompt = _load_prompt(project_root)
        from report_generator_professional import _apply_sector_profile_to_prompt
        base_prompt = _apply_sector_profile_to_prompt(base_prompt, cfg)
        model = cfg.get("openai_model", "gpt-4o-mini")
        temperature = cfg.get("openai_temperature", 0.3)
        title = cfg.get("report_title", "Informe ORM y CX")

        client_clean = _normalize_client_name(client_name)
        safe_name = _safe_filename(client_clean)
        client_dir = Path(cfg.get("output_dir", "./data/reports")) / safe_name
        client_dir.mkdir(parents=True, exist_ok=True)

        # Copy source CSV if available
        if "source_csv" in cfg and Path(cfg["source_csv"]).exists():
            try:
                shutil.copy(cfg["source_csv"], client_dir / Path(cfg["source_csv"]).name)
            except Exception as e:
                logger.warning(f"Could not copy source CSV: {e}")

        # Generate AI content
        reviews_text = _build_reviews_text(df)
        full_prompt = base_prompt.replace("{{reviews_text}}", reviews_text)
        ai_text = _call_openai(full_prompt, model, temperature)

        # Generate charts
        g1 = _rating_chart_by_star(df, client_dir / "rating_by_star.png")
        g2 = _sentiment_chart(df, client_dir / "sentiment_chart.png")

        # Create document
        doc = Document()
        _apply_doc_standards(doc, client_clean)

        # Title
        title_para = doc.add_paragraph(f"{title} — {client_clean}")
        title_para.style = doc.styles['Heading 1']
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.runs[0].font.color.rgb = PRIMARY

        # KPIs
        avg = df["rating_num"].mean() if len(df) else 0
        total = len(df)
        pos = (df["sentiment"].isin(["positivo", "muy_positivo"]).sum() / total * 100) if total else 0
        
        _kpi_card_row(doc, [
            ("Valoración media", f"{avg:.2f}", "Promedio global"),
            ("Total reseñas", f"{total}", "Volumen analizado"),
            ("% positivas", f"{pos:.0f}%", "Positivo + Muy positivo")
        ])

        # Charts
        charts_heading = doc.add_paragraph("Indicadores visuales")
        charts_heading.style = doc.styles['Heading 2']
        
        if g1 and Path(g1).exists():
            doc.add_picture(str(g1), width=Inches(5))
        if g2 and Path(g2).exists():
            doc.add_picture(str(g2), width=Inches(5))

        # Simple TOC
        doc.add_page_break()
        _insert_simple_toc(doc, "Contenido del informe")
        doc.add_page_break()

        # AI analysis
        analysis_heading = doc.add_paragraph("Análisis detallado")
        analysis_heading.style = doc.styles['Heading 2']
        
        for line in ai_text.splitlines():
            if line.strip():
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)

        # Save document
        output_file = client_dir / f"{safe_name}_informe_IA_PRO.docx"
        return _save_word_compatible(doc, output_file)

    except Exception as e:
        logger.error(f"Error in report generation: {e}")
        raise
