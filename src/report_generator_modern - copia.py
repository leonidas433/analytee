# report_generator_modern.py — Versión moderna con formato profesional estilo Excel
# ORM Analyzer IA PRO — Versión Moderna

import os, re, shutil, unicodedata, time
from pathlib import Path
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# Colores modernos
PRIMARY = RGBColor(0x00, 0x4E, 0x89)      # Azul corporativo
SECONDARY = RGBColor(0x1E, 0x88, 0xE5)    # Azul claro
SUCCESS = RGBColor(0x43, 0xA0, 0x47)      # Verde
WARNING = RGBColor(0xFB, 0x8C, 0x00)      # Naranja
DANGER = RGBColor(0xE5, 0x39, 0x35)       # Rojo
TEXT = RGBColor(0x2C, 0x2C, 0x2C)         # Gris oscuro
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)   # Gris claro
WHITE = RGBColor(0xFF, 0xFF, 0xFF)        # Blanco

# ========================== Utils ==========================
def _normalize_client_name(name: str) -> str:
    name = re.sub(r'_ChIJ[\w-]+$', '', name)
    name = name.replace('_', ' ').strip()
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("utf-8")
    name = re.sub(r'[^A-Za-z0-9 áéíóúÁÉÍÓÚüÜñÑ-]', '', name)
    return name.strip()

def _safe_filename(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9 _-]', '', name)

def _set_cell_shade(cell, hex_color):
    """Aplicar color de fondo a una celda"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shade_xml = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>'
        )
        tcPr.append(shade_xml)
    except Exception as e:
        logger.warning(f"No se pudo aplicar color de fondo: {e}")

def _set_cell_border(cell, style="single", color="000000", width="4"):
    """Aplicar bordes a una celda"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:left w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:bottom w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:right w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
    except Exception as e:
        logger.warning(f"No se pudo aplicar borde: {e}")

# ========================== Documento styling ==========================
def _apply_modern_doc_standards(doc, client_display):
    """Aplicar formato moderno al documento"""
    try:
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

            # Header moderno
            header = section.header.paragraphs[0]
            run = header.add_run(f"📊 Informe ORM — {client_display}")
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY
            run.font.bold = True
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer moderno
            footer = section.footer.paragraphs[0]
            footer_run = footer.add_run(f"📅 {datetime.now().strftime('%d/%m/%Y')} • Página [X] de [Y]")
            footer_run.font.name = "Segoe UI"
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = TEXT

        # Estilo normal moderno
        normal = doc.styles["Normal"]
        normal.font.name = "Segoe UI"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        
    except Exception as e:
        logger.warning(f"No se pudo aplicar todos los estándares: {e}")

# ========================== Tablas modernas ==========================
def _create_modern_kpi_table(doc, kpis):
    """Crear tabla de KPIs moderna estilo Excel"""
    try:
        table = doc.add_table(rows=1, cols=len(kpis))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Configurar ancho de columnas
        for i in range(len(kpis)):
            table.columns[i].width = Cm(4.5)
        
        for i, (title, value, note, color) in enumerate(kpis):
            cell = table.cell(0, i)
            
            # Aplicar color de fondo
            _set_cell_shade(cell, color)
            _set_cell_border(cell, "single", "FFFFFF", "8")
            
            # Configurar alineación vertical
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Título
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(title.upper())
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = PRIMARY
            r1.font.name = "Segoe UI"
            
            # Valor principal
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(str(value))
            r2.bold = True
            r2.font.size = Pt(24)
            r2.font.color.rgb = TEXT
            r2.font.name = "Segoe UI"
            
            # Nota
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run(note)
            r3.font.size = Pt(9)
            r3.font.color.rgb = TEXT
            r3.font.name = "Segoe UI"
            r3.italic = True
            
    except Exception as e:
        logger.error(f"Error creando tabla KPI: {e}")

def _create_data_summary_table(doc, df):
    """Crear tabla resumen de datos moderna"""
    try:
        # Calcular estadísticas
        total_reviews = len(df)
        avg_rating = df['rating_num'].mean() if total_reviews > 0 else 0
        
        # Distribución por estrellas
        star_dist = df['rating_num'].value_counts().sort_index()
        star_5 = star_dist.get(5, 0)
        star_4 = star_dist.get(4, 0)
        star_3 = star_dist.get(3, 0)
        star_2 = star_dist.get(2, 0)
        star_1 = star_dist.get(1, 0)
        
        # Distribución por sentimiento
        sentiment_dist = df['sentiment'].value_counts()
        muy_positivo = sentiment_dist.get('muy_positivo', 0)
        positivo = sentiment_dist.get('positivo', 0)
        neutro = sentiment_dist.get('neutro', 0)
        negativo = sentiment_dist.get('negativo', 0)
        muy_negativo = sentiment_dist.get('muy_negativo', 0)
        
        # Crear tabla
        table = doc.add_table(rows=7, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Configurar anchos
        table.columns[0].width = Cm(6)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        
        # Encabezados
        headers = [
            ("Métrica", "Valor", "Porcentaje"),
            ("Total de reseñas", f"{total_reviews:,}", "100%"),
            ("Valoración promedio", f"{avg_rating:.2f} ⭐", f"{(avg_rating/5*100):.1f}%"),
            ("Reseñas 5 estrellas", f"{star_5:,}", f"{(star_5/total_reviews*100):.1f}%" if total_reviews > 0 else "0%"),
            ("Reseñas 4 estrellas", f"{star_4:,}", f"{(star_4/total_reviews*100):.1f}%" if total_reviews > 0 else "0%"),
            ("Reseñas 3 estrellas", f"{star_3:,}", f"{(star_3/total_reviews*100):.1f}%" if total_reviews > 0 else "0%"),
            ("Reseñas 2 estrellas", f"{star_2:,}", f"{(star_2/total_reviews*100):.1f}%" if total_reviews > 0 else "0%"),
            ("Reseñas 1 estrella", f"{star_1:,}", f"{(star_1/total_reviews*100):.1f}%" if total_reviews > 0 else "0%")
        ]
        
        for row_idx, (metric, value, percentage) in enumerate(headers):
            for col_idx in range(3):
                cell = table.cell(row_idx, col_idx)
                
                # Color de fondo para encabezado
                if row_idx == 0:
                    _set_cell_shade(cell, "E7F0FA")
                    _set_cell_border(cell, "single", "004E89", "12")
                else:
                    _set_cell_border(cell, "single", "CCCCCC", "4")
                
                # Configurar alineación
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                
                # Texto
                run = p.add_run(metric if col_idx == 0 else (value if col_idx == 1 else percentage))
                run.font.name = "Segoe UI"
                run.font.size = Pt(10)
                
                if row_idx == 0:  # Encabezado
                    run.bold = True
                    run.font.color.rgb = PRIMARY
                else:
                    run.font.color.rgb = TEXT
                    
    except Exception as e:
        logger.error(f"Error creando tabla resumen: {e}")

def _create_languages_table_modern(doc, languages):
    """Crear tabla de distribución de idiomas moderna"""
    try:
        if not languages:
            return

        # Crear tabla
        table = doc.add_table(rows=len(languages) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Configurar anchos
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(2.5)
        table.columns[2].width = Cm(2.5)
        table.columns[3].width = Cm(6)

        # Encabezado
        header_cells = ["Idioma", "Cantidad", "Porcentaje", "Análisis Estratégico"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY
            run.font.name = "Segoe UI"

        # Filas de datos
        for row_idx, lang in enumerate(languages, 1):
            for col_idx in range(4):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in [0, 3] else WD_ALIGN_PARAGRAPH.CENTER

                if col_idx == 0:
                    content = lang["name"]
                    run = p.add_run(content)
                    run.bold = True
                    run.font.color.rgb = PRIMARY
                elif col_idx == 1:
                    content = f"{lang['count']:,}"
                    run = p.add_run(content)
                    run.bold = True
                elif col_idx == 2:
                    content = f"{lang['percentage']:.1f}%"
                    run = p.add_run(content)
                    run.bold = True
                elif col_idx == 3:
                    # Análisis estratégico
                    if lang['percentage'] > 50:
                        analisis = "🌍 Mercado principal - Comunicación prioritaria"
                        color = SUCCESS
                    elif lang['percentage'] > 20:
                        analisis = "🎯 Mercado secundario - Atención especializada"
                        color = WARNING
                    elif lang['percentage'] > 10:
                        analisis = "📈 Mercado emergente - Oportunidad de crecimiento"
                        color = SECONDARY
                    else:
                        analisis = "🔍 Mercado nicho - Monitoreo selectivo"
                        color = TEXT
                    run = p.add_run(analisis)
                    run.font.color.rgb = color

                run.font.name = "Segoe UI"
                run.font.size = Pt(9)
                if col_idx < 3:
                    run.font.color.rgb = TEXT

    except Exception as e:
        logger.error(f"Error creando tabla idiomas moderna: {e}")


def _create_sentiment_analysis_table(doc, df):
    """Crear tabla de análisis de sentimiento moderna"""
    try:
        sentiment_dist = df['sentiment'].value_counts()
        total = len(df)

        # Datos ordenados
        sentiments = [
            ("Muy Positivo", "muy_positivo", SUCCESS),
            ("Positivo", "positivo", SUCCESS),
            ("Neutro", "neutro", WARNING),
            ("Negativo", "negativo", DANGER),
            ("Muy Negativo", "muy_negativo", DANGER)
        ]

        table = doc.add_table(rows=len(sentiments) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Configurar anchos
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(3)
        table.columns[3].width = Cm(4)

        # Encabezado
        header_cells = ["Sentimiento", "Cantidad", "Porcentaje", "Barra Visual"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY
            run.font.name = "Segoe UI"

        # Filas de datos
        for row_idx, (label, key, color) in enumerate(sentiments, 1):
            count = sentiment_dist.get(key, 0)
            percentage = (count / total * 100) if total > 0 else 0

            # Sentimiento
            cell = table.cell(row_idx, 0)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(label)
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            run.font.color.rgb = color

            # Cantidad
            cell = table.cell(row_idx, 1)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{count:,}")
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            run.bold = True

            # Porcentaje
            cell = table.cell(row_idx, 2)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{percentage:.1f}%")
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            run.bold = True

            # Barra visual (usando caracteres)
            cell = table.cell(row_idx, 3)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bar_length = int(percentage / 2)  # Escala visual
            bar = "█" * bar_length + "░" * (20 - bar_length)
            run = p.add_run(bar)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = color

    except Exception as e:
        logger.error(f"Error creando tabla sentimiento: {e}")

# ========================== Gráficos modernos ==========================
def _create_modern_rating_chart(df, path):
    """Crear gráfico de valoraciones moderno"""
    try:
        counts = df['rating_num'].value_counts().sort_index()
        total = float(counts.sum()) if len(counts) else 0.0
        
        # Configurar estilo moderno
        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Colores modernos
        colors = ['#E53935', '#FF7043', '#FFB74D', '#66BB6A', '#43A047']
        bars = ax.bar(counts.index.astype(str), counts.values, color=colors, 
                     edgecolor='white', linewidth=2, alpha=0.9)
        
        # Personalizar gráfico
        ax.set_xlabel('Valoración (Estrellas)', fontsize=12, fontweight='bold', color='#2C2C2C')
        ax.set_ylabel('Número de Reseñas', fontsize=12, fontweight='bold', color='#2C2C2C')
        ax.set_title('Distribución de Valoraciones por Estrellas', 
                    fontsize=14, fontweight='bold', color='#004E89', pad=20)
        
        # Añadir valores en las barras
        for bar in bars:
            height = bar.get_height()
            pct = (float(height) / total * 100.0) if total else 0.0
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                   f'{f"{int(height):,}".replace(",", ".")} ({pct:.1f}%)', ha='center', va='bottom', 
                   fontweight='bold', fontsize=10)
        
        # Estilo de la grilla
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        ax.set_axisbelow(True)
        
        # Colores de ejes
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#CCCCCC')
        ax.spines['bottom'].set_color('#CCCCCC')
        
        plt.tight_layout()
        plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        return path
        
    except Exception as e:
        logger.error(f"Error creando gráfico de valoraciones: {e}")
        return None

def _create_modern_sentiment_chart(df, path):
    """Crear gráfico de sentimiento moderno"""
    try:
        order = ["muy_negativo", "negativo", "neutro", "positivo", "muy_positivo"]
        labels = ["Muy Negativo", "Negativo", "Neutro", "Positivo", "Muy Positivo"]
        colors = ["#C62828", "#E53935", "#FB8C00", "#43A047", "#2E7D32"]

        v = df['sentiment'].value_counts()
        vals = [int(v.get(k, 0)) for k in order]
        total = float(sum(vals)) if len(vals) else 0.0

        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(10, 6))

        bars = ax.barh(labels, vals, color=colors, edgecolor='white', linewidth=2, alpha=0.9)

        # Personalizar gráfico
        ax.set_xlabel('Número de Reseñas', fontsize=12, fontweight='bold', color='#2C2C2C')
        ax.set_ylabel('Sentimiento', fontsize=12, fontweight='bold', color='#2C2C2C')
        ax.set_title('Distribución por Análisis de Sentimiento',
                    fontsize=14, fontweight='bold', color='#004E89', pad=20)

        # Añadir valores en las barras
        for i, bar in enumerate(bars):
            width = bar.get_width()
            pct = (float(vals[i]) / total * 100.0) if total else 0.0
            ax.text(width + 0.1, bar.get_y() + bar.get_height()/2.,
                   f'{f"{int(vals[i]):,}".replace(",", ".")} ({pct:.1f}%)', ha='left', va='center',
                   fontweight='bold', fontsize=10)

        # Estilo de la grilla
        ax.grid(axis='x', alpha=0.3, linestyle='--')
        ax.set_axisbelow(True)

        # Colores de ejes
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#CCCCCC')
        ax.spines['bottom'].set_color('#CCCCCC')

        plt.tight_layout()
        plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        return path

    except Exception as e:
        logger.error(f"Error creando gráfico de sentimiento: {e}")
        return None


def _create_modern_language_chart(languages, path):
    """Crear gráfico de distribución de idiomas moderno"""
    try:
        if not languages:
            return None

        # Preparar datos
        labels = [lang['name'] for lang in languages]
        sizes = [lang['percentage'] for lang in languages]
        counts = [int(lang.get("count", 0)) for lang in languages]
        colors = ['#004E89', '#1E88E5', '#43A047', '#FB8C00', '#E53935', '#8E24AA', '#3949AB', '#00ACC1', '#7CB342', '#F4511E']
        legend_labels = [
            f"{lang['name']}: {f'{int(lang.get('count', 0)):,}'.replace(',', '.')} ({lang.get('percentage', 0)}%)"
            for lang in languages
        ]

        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(12, 8))

        if len(labels) <= 6:
            idx = {"i": 0}

            def _autopct(pct):
                i = idx["i"]
                idx["i"] += 1
                v = counts[i] if i < len(counts) else 0
                return f"{f'{int(v):,}'.replace(',', '.')}\n({pct:.1f}%)"

            wedges, texts, autotexts = ax.pie(
                sizes,
                colors=colors[:len(labels)],
                autopct=_autopct,
                pctdistance=0.78,
                startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 3, 'alpha': 0.9},
            )
        else:
            wedges, texts, autotexts = ax.pie(
                sizes,
                colors=colors[:len(labels)],
                autopct=None,
                startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 3, 'alpha': 0.9},
            )

        # Personalizar
        ax.set_title('Distribución de Idiomas en Reseñas\nAnálisis Multilingüe del Cliente',
                    fontsize=16, fontweight='bold', color='#004E89', pad=20)

        # Mejorar legibilidad
        for text in texts:
            text.set_fontsize(11)
            text.set_fontweight('bold')
            text.set_color('#2C2C2C')
        for wedge, autotext in zip(wedges, autotexts):
            autotext.set_fontsize(10)
            autotext.set_fontweight('bold')
            try:
                r, g, b, _ = mcolors.to_rgba(wedge.get_facecolor())
                lum = 0.299 * r + 0.587 * g + 0.114 * b
                autotext.set_color('black' if lum >= 0.62 else 'white')
            except Exception:
                autotext.set_color('white')

        ax.legend(wedges, legend_labels, title="Distribución por Idiomas",
                 loc="center left", bbox_to_anchor=(1, 0, 0.5, 1), fontsize=10)

        plt.tight_layout()
        plt.savefig(path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        return path

    except Exception as e:
        logger.error(f"Error creando gráfico de idiomas moderno: {e}")
        return None

# ========================== Prompt helpers ==========================
def _load_prompt(project_root: Path) -> str:
    p = project_root / "src" / "prompt.txt"
    try:
        return p.read_text(encoding="utf-8-sig") if p.exists() else "{{reviews_text}}"
    except Exception as e:
        logger.warning(f"No se pudo cargar prompt: {e}")
        return "{{reviews_text}}"

def _build_reviews_text(df):
    rows = []
    for _, r in df.iterrows():
        try:
            user = str(r.get("user", "")).strip() or "Usuario"
            date = str(r.get("date", "")).strip() or "s/f"
            rating = int(r.get("rating_num", 0))
            txt = str(r.get("text", "")).strip() or "[Sin comentario]"
            rows.append(f"- [{date}] {user} ({rating}★): {txt}")
        except Exception as e:
            logger.warning(f"Error procesando fila de reseña: {e}")
            continue
    return "\n".join(rows)

def _call_openai(prompt: str, model: str, temperature: float):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return "Informe generado sin conexión a IA (defina OPENAI_API_KEY)."
    
    try:
        client = OpenAI(api_key=api_key)
        r = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[{"role": "user", "content": prompt}],
        )
        return r.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Error API OpenAI: {e}")
        return f"Error al generar contenido con IA: {e}"

def _insert_ai_content_modern(doc, ai_text):
    """
    Parser inteligente de contenido IA para Word Moderno.
    Convierte Markdown (títulos, listas) en estilos Word y limpia tablas/basura.
    """
    import re
    
    # Limpiar tablas Markdown (líneas que contienen | y ---)
    lines = ai_text.splitlines()
    clean_lines = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        # Detectar inicio/fin de tabla Markdown
        if stripped.startswith('|') or (stripped.startswith('+-') and '-+' in stripped):
            in_table = True
            continue
        if in_table and not stripped.startswith('|') and not stripped.startswith('+-'):
            in_table = False
        
        if not in_table and not re.match(r'^[|\-\s:+]{3,}$', stripped):
            clean_lines.append(line)
    
    for line in clean_lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # 1. Títulos (## o ###)
        if stripped.startswith('###'):
            title = stripped.replace('###', '').strip()
            p = doc.add_paragraph(title)
            p.style = doc.styles['Heading 3']
            # Color corporativo para títulos
            for run in p.runs:
                run.font.color.rgb = PRIMARY
            continue
            
        if stripped.startswith('##') or (re.match(r'^\d+\.\s+\*\*', stripped)):
            title = re.sub(r'^(##|\d+\.\s+)', '', stripped).replace('**', '').strip()
            p = doc.add_paragraph(title)
            p.style = doc.styles['Heading 2']
            for run in p.runs:
                run.font.color.rgb = PRIMARY
            continue

        # 2. Listas con guiones o asteriscos
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.0)
            continue

        # 3. Párrafos normales (limpiar negritas Markdown **text**)
        clean_text = stripped.replace('**', '')
        p = doc.add_paragraph(clean_text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Cm(0.5)
        for run in p.runs:
            run.font.name = "Segoe UI"
            run.font.size = Pt(11)


# ========================== Generador principal ==========================
def create_modern_report(df, client_name, output_path, cfg, project_root="."):
    """Función principal de generación de reporte moderno"""
    try:
        project_root = Path(project_root)
        base_prompt = _load_prompt(project_root)
        model = cfg.get("openai_model", "gpt-4o-mini")
        temperature = cfg.get("openai_temperature", 0.3)
        title = cfg.get("report_title", "Informe ORM y CX")

        client_clean = _normalize_client_name(client_name)
        safe_name = _safe_filename(client_clean)
        client_dir = Path(cfg.get("output_dir", "./data/reports")) / safe_name
        client_dir.mkdir(parents=True, exist_ok=True)

        # Copiar CSV fuente si está disponible
        if "source_csv" in cfg and Path(cfg["source_csv"]).exists():
            try:
                shutil.copy(cfg["source_csv"], client_dir / Path(cfg["source_csv"]).name)
            except Exception as e:
                logger.warning(f"No se pudo copiar CSV fuente: {e}")

        # Generar contenido IA
        reviews_text = _build_reviews_text(df)
        full_prompt = base_prompt.replace("{{reviews_text}}", reviews_text)
        ai_text = _call_openai(full_prompt, model, temperature)

        # Generar gráficos modernos
        g1 = _create_modern_rating_chart(df, client_dir / "rating_by_star_modern.png")
        g2 = _create_modern_sentiment_chart(df, client_dir / "sentiment_chart_modern.png")
        g3 = _create_modern_language_chart(cfg.get("languages", []), client_dir / "language_chart_modern.png")

        # Crear documento
        doc = Document()
        _apply_modern_doc_standards(doc, client_clean)

        # Portada moderna
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"📊 {title}")
        title_run.font.name = "Segoe UI"
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = PRIMARY
        title_run.bold = True

        client_para = doc.add_paragraph()
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client_para.add_run(f"Cliente: {client_clean}")
        client_run.font.name = "Segoe UI"
        client_run.font.size = Pt(18)
        client_run.font.color.rgb = TEXT

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%d de %B de %Y"))
        date_run.font.name = "Segoe UI"
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = TEXT
        date_run.italic = True

        doc.add_page_break()

        table_no = 1

        # KPIs modernos
        avg = df["rating_num"].mean() if len(df) else 0
        total = len(df)
        pos = (df["sentiment"].isin(["positivo", "muy_positivo"]).sum() / total * 100) if total else 0
        neg = (df["sentiment"].isin(["negativo", "muy_negativo"]).sum() / total * 100) if total else 0
        
        kpi_data = [
            ("Valoración Media", f"{avg:.2f} ⭐", "Promedio global", "E7F0FA"),
            ("Total Reseñas", f"{total:,}", "Volumen analizado", "E8F5E8"),
            ("% Positivas", f"{pos:.0f}%", "Positivo + Muy positivo", "E8F5E8"),
            ("% Negativas", f"{neg:.0f}%", "Negativo + Muy negativo", "FFEBEE")
        ]
        
        doc.add_paragraph(f"Tabla {table_no} – Indicadores clave de reputación (KPIs)")
        table_no += 1
        _create_modern_kpi_table(doc, kpi_data)

        # Espaciado
        doc.add_paragraph()

        # Tabla resumen de datos
        summary_heading = doc.add_paragraph("📈 Resumen de Datos")
        summary_heading.style = doc.styles['Heading 2']
        summary_heading.runs[0].font.color.rgb = PRIMARY
        
        doc.add_paragraph(f"Tabla {table_no} – Resumen de datos y distribución por estrellas")
        table_no += 1
        _create_data_summary_table(doc, df)

        # Espaciado
        doc.add_paragraph()

        # Gráficos modernos
        charts_heading = doc.add_paragraph("📊 Indicadores Visuales")
        charts_heading.style = doc.styles['Heading 2']
        charts_heading.runs[0].font.color.rgb = PRIMARY
        
        if g1 and Path(g1).exists():
            doc.add_picture(str(g1), width=Inches(6))
        if g2 and Path(g2).exists():
            doc.add_picture(str(g2), width=Inches(6))

        # Espaciado
        doc.add_paragraph()

        # Tabla de análisis de sentimiento
        sentiment_heading = doc.add_paragraph("🎯 Análisis de Sentimiento Detallado")
        sentiment_heading.style = doc.styles['Heading 2']
        sentiment_heading.runs[0].font.color.rgb = PRIMARY

        doc.add_paragraph(f"Tabla {table_no} – Distribución de sentimiento de las reseñas")
        table_no += 1
        _create_sentiment_analysis_table(doc, df)

        # Espaciado
        doc.add_paragraph()

        # Tabla de distribución de idiomas
        languages_heading = doc.add_paragraph("🌍 Análisis de Idiomas")
        languages_heading.style = doc.styles['Heading 2']
        languages_heading.runs[0].font.color.rgb = PRIMARY

        languages = cfg.get("languages", [])
        if languages:
            doc.add_paragraph(f"Tabla {table_no} – Distribución de idiomas de las reseñas")
            table_no += 1
        _create_languages_table_modern(doc, languages)

        # Espaciado
        doc.add_paragraph()

        # Gráfico de idiomas
        if g3 and Path(g3).exists():
            doc.add_picture(str(g3), width=Inches(6))

        # Salto de página
        doc.add_page_break()

        # Análisis IA
        analysis_heading = doc.add_paragraph("🤖 Análisis Estratégico e Insights")
        analysis_heading.style = doc.styles['Heading 2']
        analysis_heading.runs[0].font.color.rgb = PRIMARY
        
        # Insertar contenido IA procesado
        _insert_ai_content_modern(doc, ai_text)

        # Guardar documento
        output_file = client_dir / f"{safe_name}_informe_MODERNO.docx"
        doc.save(str(output_file))
        logger.info(f"✅ Informe moderno guardado: {output_file}")
        
        # Intentar generar PDF
        try:
            from docx2pdf import convert
            pdf_path = output_file.with_suffix(".pdf")
            convert(str(output_file), str(pdf_path))
            logger.info(f"📄 PDF generado: {pdf_path}")
        except Exception as e:
            logger.warning(f"No se pudo generar PDF: {e}")

        return str(output_file)

    except Exception as e:
        logger.error(f"Error en generación de reporte: {e}")
        raise
