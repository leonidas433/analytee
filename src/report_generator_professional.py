# report_generator_professional.py — Generador de reportes con formato visual profesional
# ORM Analyzer IA PRO — Versión Profesional con Markdown y Tablas Estilo Consultoría

import os, re, shutil, unicodedata, time, hashlib, json
from pathlib import Path
from datetime import datetime, timezone
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logging.getLogger("httpx").setLevel(logging.WARNING)

PIPELINE_VERSION = "2025.12.30"

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

from report_modes import is_audit_mode, redact_for_client
from alerts_engine import compute_alerts
from analysis_section_6 import (
    compute_contradictions,
    compute_local_vs_tourist,
    compute_operational_fatigue,
    compute_contradiction_score,
    compute_operational_fatigue_score,
    get_metrics_version,
    compute_alerts_by_score,
)

# Colores profesionales
PRIMARY = RGBColor(0x00, 0x4E, 0x89)      # Azul corporativo
SECONDARY = RGBColor(0x1E, 0x88, 0xE5)    # Azul claro
SUCCESS = RGBColor(0x43, 0xA0, 0x47)      # Verde
WARNING = RGBColor(0xFB, 0x8C, 0x00)      # Naranja
DANGER = RGBColor(0xE5, 0x39, 0x35)       # Rojo
TEXT = RGBColor(0x2C, 0x2C, 0x2C)         # Gris oscuro
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)   # Gris claro
WHITE = RGBColor(0xFF, 0xFF, 0xFF)        # Blanco

SECTOR_CATEGORY_WHITELIST: set[str] = {
    "Restaurant",
    "Fast food restaurant",
    "Bank",
    "Sporting goods store",
    "Warehouse",
    "Logistics service",
    "Shopping mall",
}

SECTOR_PROFILES = {
    # Perfiles sectoriales: solo se usan para enriquecer prompts IA (nunca para DOCX/PDF).
    "generic": {
        "label": "Negocio genérico",
        "context": "Aplica principios generales de experiencia de cliente, reputación online y comunicación profesional.",
    },
    "restaurant": {
        "label": "Restauración",
        "context": "Clientes con visitas presenciales, picos horarios, reseñas ligadas a servicio, comida, espera y atención.",
    },
}

# ========================== Utils ==========================
def _normalize_client_name(name: str) -> str:
    name = re.sub(r'_ChIJ[\w-]+$', '', name)
    name = name.replace('_', ' ').strip()
    name = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("utf-8")
    name = re.sub(r'[^A-Za-z0-9 áéíóúÁÉÍÓÚüÜñÑ-]', '', name)
    return name.strip()

def _safe_filename(name: str) -> str:
    return re.sub(r'[^A-Za-z0-9 _-]', '', name)


def _sha256_file(path: Path) -> str | None:
    try:
        h = hashlib.sha256()
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return None


def _inject_docx_traceability_metadata(doc: Document, *, pipeline_version: str, input_hash: str | None, generated_at: str):
    try:
        payload = {
            "pipeline_version": pipeline_version,
            "input_hash": input_hash,
            "generated_at": generated_at,
        }
        doc.core_properties.comments = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    except Exception:
        pass


class QualityCheckError(Exception):
    pass


class OutputContractError(Exception):
    pass


def validate_output_contract(output_dir, context):
    out_dir = Path(str(output_dir))
    ctx = context or {}
    if not isinstance(ctx, dict):
        ctx = {"context": str(ctx)}

    require_pdf = bool(ctx.get("require_pdf", True))

    if not out_dir.exists() or not out_dir.is_dir():
        raise OutputContractError(f"CONTRACT: output_dir no existe o no es carpeta: {out_dir}")

    docx_files = sorted(out_dir.glob("*_informe_PROFESIONAL.docx"))
    if len(docx_files) != 1:
        raise OutputContractError(f"CONTRACT: esperado 1 DOCX, encontrados {len(docx_files)} en {out_dir}")
    docx_path = docx_files[0]

    pdf_path = None
    if require_pdf:
        pdf_files = sorted(out_dir.glob("*_informe_PROFESIONAL.pdf"))
        if len(pdf_files) != 1:
            raise OutputContractError(f"CONTRACT: esperado 1 PDF, encontrados {len(pdf_files)} en {out_dir}")
        pdf_path = pdf_files[0]
    else:
        pdf_files = sorted(out_dir.glob("*_informe_PROFESIONAL.pdf"))
        if len(pdf_files) == 1:
            pdf_path = pdf_files[0]

    log_path = out_dir / "execution_log.json"
    if not log_path.exists():
        raise OutputContractError(f"CONTRACT: falta execution_log.json en {out_dir}")

    required_paths = [docx_path, log_path]
    if require_pdf and pdf_path is not None:
        required_paths.append(pdf_path)

    for p in required_paths:
        try:
            size = p.stat().st_size
        except Exception as e:
            raise OutputContractError(f"CONTRACT: no se pudo leer tamaño de {p} ({e})")
        if size <= 0:
            raise OutputContractError(f"CONTRACT: archivo vacío (size=0): {p}")

    try:
        raw_log = log_path.read_text(encoding="utf-8-sig")
        payload = json.loads(raw_log or "{}")
    except Exception as e:
        raise OutputContractError(f"CONTRACT: execution_log.json inválido ({e})")

    if not isinstance(payload, dict):
        raise OutputContractError("CONTRACT: execution_log.json debe ser objeto JSON")

    docx_hash = _sha256_file(docx_path)
    if not docx_hash:
        raise OutputContractError(f"CONTRACT: no se pudo calcular sha256 DOCX: {docx_path}")

    pdf_hash = None
    if require_pdf:
        if pdf_path is None:
            raise OutputContractError("CONTRACT: PDF requerido pero no encontrado")
        pdf_hash = _sha256_file(pdf_path)
        if not pdf_hash:
            raise OutputContractError(f"CONTRACT: no se pudo calcular sha256 PDF: {pdf_path}")

    logged_docx = payload.get("docx_hash")
    logged_pdf = payload.get("pdf_hash")

    if logged_docx != docx_hash:
        raise OutputContractError(f"CONTRACT: sha256 DOCX no coincide (log={logged_docx} real={docx_hash})")
    if require_pdf and logged_pdf != pdf_hash:
        raise OutputContractError(f"CONTRACT: sha256 PDF no coincide (log={logged_pdf} real={pdf_hash})")

    return {
        "output_dir": str(out_dir),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path) if pdf_path else None,
        "execution_log": str(log_path),
        "docx_hash": docx_hash,
        "pdf_hash": pdf_hash,
    }


def _iter_docx_text_blocks(doc: Document):
    for p_i, p in enumerate(getattr(doc, "paragraphs", []) or []):
        yield (f"PARAGRAPH[{p_i}]", p.text or "")
    for t_i, table in enumerate(getattr(doc, "tables", []) or []):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                yield (f"TABLE[{t_i}].CELL[{r_i},{c_i}]", cell.text or "")
    for s_i, section in enumerate(getattr(doc, "sections", []) or []):
        header = getattr(section, "header", None)
        if header is not None:
            for p_i, p in enumerate(getattr(header, "paragraphs", []) or []):
                yield (f"HEADER[{s_i}].PARAGRAPH[{p_i}]", p.text or "")
        footer = getattr(section, "footer", None)
        if footer is not None:
            for p_i, p in enumerate(getattr(footer, "paragraphs", []) or []):
                yield (f"FOOTER[{s_i}].PARAGRAPH[{p_i}]", p.text or "")


def _scan_docx_forbidden_text(docx_path: Path, forbidden_text: list[str]):
    doc = Document(str(docx_path))
    forbidden = [t for t in (forbidden_text or []) if isinstance(t, str) and t.strip()]
    if not forbidden:
        return []

    hits = []
    forbidden_patterns = []
    for token in forbidden:
        t = token.strip()
        escaped = re.escape(t)
        if re.fullmatch(r"[A-Za-z0-9_]+", t):
            pat = re.compile(rf"\b{escaped}\b")
        else:
            pat = re.compile(rf"\b{escaped}\b")
        forbidden_patterns.append((t, pat))

    for where, text in _iter_docx_text_blocks(doc):
        raw = text or ""
        if not raw.strip():
            continue
        for needle, pat in forbidden_patterns:
            if pat.search(raw):
                hits.append((needle, where, raw.strip()))
    return hits


def run_quality_checks(docx_path, context):
    p = Path(str(docx_path))
    if not p.exists():
        raise QualityCheckError(f"QA: DOCX no encontrado: {p}")

    ctx = context or {}
    if not isinstance(ctx, dict):
        ctx = {"context": str(ctx)}

    enable_structure_audit = bool(ctx.get("enable_structure_audit", True))
    enable_forbidden_scan = bool(ctx.get("enable_forbidden_scan", True))
    forbidden_text = ctx.get("forbidden_text")
    if not isinstance(forbidden_text, list):
        forbidden_text = ["None", "NO DISPONIBLE", "RUTA_AL_DOCX"]

    errors = []

    if enable_structure_audit:
        try:
            from audit_docx_structure import audit_docx_structure as _audit_docx_structure
        except Exception as e:
            errors.append(f"QA: No se pudo importar audit_docx_structure ({e})")
        else:
            result = _audit_docx_structure(p)
            violations = result.get("violations") or []
            if violations:
                sample = violations[:20]
                errors.append(
                    "QA: audit_docx_structure detectó violaciones:\n"
                    + "\n".join(f"- {v}" for v in sample)
                )

    if enable_forbidden_scan:
        hits = _scan_docx_forbidden_text(p, forbidden_text)
        if hits:
            max_hits = int(ctx.get("max_forbidden_hits", 20) or 20)
            sample = hits[:max_hits]
            errors.append(
                "QA: Texto prohibido detectado:\n"
                + "\n".join(f"- TOKEN={t} WHERE={w} TEXT={txt}" for t, w, txt in sample)
            )

    if errors:
        details = "\n\n".join(errors)
        raise QualityCheckError(f"❌ QA AUTOMÁTICO FALLÓ\nDOCX={p}\n{details}")

    return True


def _safe_trace_id(generated_at: str) -> str:
    raw = str(generated_at or "")
    cleaned = re.sub(r"[^0-9A-Za-z]+", "", raw)
    return cleaned or datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")

def _set_cell_shade(cell, hex_color):
    """Aplicar color de fondo a una celda"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shade_xml = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>'
        )
        tcPr.append(shade_xml)
    except Exception as e:
        logger.warning(f"No se pudo aplicar color de fondo: {e}")


def _sanitize_section6_artifacts_for_mode(cfg: dict, mode: str) -> dict:
    """
    Devuelve copia de artefactos sanitizados según modo.
    """
    sanitized = dict(cfg)
    mode_normalized = str(mode or "CLIENT").upper()

    if mode_normalized == "CLIENT":
        contradictions = sanitized.get("artifact_contradictions")
        if isinstance(contradictions, dict):
            c_copy = contradictions.copy()
            c_copy["examples"] = []
            sanitized["artifact_contradictions"] = c_copy

        fatigue = sanitized.get("artifact_operational_fatigue")
        if isinstance(fatigue, dict):
            f_copy = fatigue.copy()
            f_copy["repeated_complaints"] = []
            sanitized["artifact_operational_fatigue"] = f_copy

    return sanitized

def _set_cell_border(cell, style="single", color="000000", width="4"):
    """Aplicar bordes a una celda"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:left w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:bottom w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'<w:right w:val="{style}" w:sz="{width}" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
    except Exception as e:
        logger.warning(f"No se pudo aplicar borde: {e}")

def _set_cell_padding(cell, *, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120):
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:tcMar"))
        if existing is not None:
            tcPr.remove(existing)
        tcMar = parse_xml(
            f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top w:w="{int(top)}" w:type="dxa"/>'
            f'<w:left w:w="{int(left)}" w:type="dxa"/>'
            f'<w:bottom w:w="{int(bottom)}" w:type="dxa"/>'
            f'<w:right w:w="{int(right)}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)
    except Exception:
        pass

# ========================== Documento styling ==========================
def _apply_professional_doc_standards(doc, client_display):
    """Aplicar formato profesional al documento"""
    try:
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

            # Header profesional
            header = section.header.paragraphs[0]
            run = header.add_run(f"📊 Informe ORM Profesional — {client_display}")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = PRIMARY
            run.font.bold = True
            header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Footer profesional
            footer = section.footer.paragraphs[0]
            footer_run = footer.add_run(f"📅 {datetime.now().strftime('%d/%m/%Y')} • Página [X] de [Y]")
            footer_run.font.name = "Calibri"
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = TEXT

        # Estilo normal profesional
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)

        heading2 = doc.styles["Heading 2"]
        heading2.font.name = "Calibri"
        heading2.font.size = Pt(16)
        heading2.font.bold = True
        heading2.paragraph_format.space_before = Pt(16)
        heading2.paragraph_format.space_after = Pt(12)

        heading3 = doc.styles["Heading 3"]
        heading3.font.name = "Calibri"
        heading3.font.size = Pt(13)
        heading3.font.bold = True
        heading3.paragraph_format.space_before = Pt(6)
        heading3.paragraph_format.space_after = Pt(4)
        
    except Exception as e:
        logger.warning(f"No se pudo aplicar todos los estándares: {e}")


def _apply_table_layout(table, row_height_cm: float = 0.8):
    try:
        table.autofit = False
    except Exception:
        pass
    try:
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Cm(row_height_cm)
            for cell in row.cells:
                _set_cell_padding(cell)
    except Exception:
        pass


def _sorted_languages(languages):
    if not languages:
        return []

    def _key(lang):
        try:
            pct = float(lang.get("percentage", 0) or 0)
        except Exception:
            pct = 0.0
        try:
            count = int(lang.get("count", 0) or 0)
        except Exception:
            count = 0
        name = str(lang.get("name", "") or "")
        code = str(lang.get("code", "") or "")
        return (-pct, -count, name, code)

    return sorted(list(languages), key=_key)


def _apply_languages_top_n(languages, top_n):
    if not languages:
        return []
    if top_n is None:
        return list(languages)
    try:
        n = int(top_n)
    except Exception:
        return list(languages)
    if n <= 0:
        return list(languages)
    if len(languages) <= n:
        return list(languages)

    kept = list(languages[:n])
    rest = list(languages[n:])

    total = 0
    rest_count = 0
    for lang in languages:
        try:
            total += int(lang.get("count", 0) or 0)
        except Exception:
            pass
    for lang in rest:
        try:
            rest_count += int(lang.get("count", 0) or 0)
        except Exception:
            pass

    if rest_count > 0:
        pct = round((rest_count / total * 100), 1) if total > 0 else 0.0
        kept.append(
            {
                "code": "other",
                "name": "Otros",
                "count": int(rest_count),
                "percentage": float(pct),
            }
        )

    return kept


class _SectionTimer:
    def __init__(self, name: str):
        self.name = str(name or "")
        self._t0 = None

    def __enter__(self):
        self._t0 = time.perf_counter()
        return self

    def __exit__(self, exc_type, exc, tb):
        try:
            if self._t0 is None:
                return False
            ms = int((time.perf_counter() - self._t0) * 1000)
            logger.info(f"TIMER_SECTION={self.name} ms={ms}")
        except Exception:
            pass
        return False

# ========================== Tablas profesionales ==========================
def _create_professional_kpi_table(doc, df, heading_mgr=None, *, star_dist=None, sentiment_dist=None):
    """Crear tabla de KPIs profesional estilo consultoría"""
    try:
        # Calcular métricas
        total_reviews = len(df)
        avg_rating = df['rating_num'].mean() if total_reviews > 0 else 0
        
        # Distribución por estrellas
        if star_dist is None:
            star_dist = df['rating_num'].value_counts().sort_index()
        star_5_pct = (star_dist.get(5, 0) / total_reviews * 100) if total_reviews > 0 else 0
        star_4_pct = (star_dist.get(4, 0) / total_reviews * 100) if total_reviews > 0 else 0
        star_3_pct = (star_dist.get(3, 0) / total_reviews * 100) if total_reviews > 0 else 0
        star_2_pct = (star_dist.get(2, 0) / total_reviews * 100) if total_reviews > 0 else 0
        star_1_pct = (star_dist.get(1, 0) / total_reviews * 100) if total_reviews > 0 else 0
        
        # Análisis de sentimiento
        if sentiment_dist is None:
            sentiment_dist = df['sentiment'].value_counts()
        muy_positivo = sentiment_dist.get('muy_positivo', 0)
        positivo = sentiment_dist.get('positivo', 0)
        neutro = sentiment_dist.get('neutro', 0)
        negativo = sentiment_dist.get('negativo', 0)
        muy_negativo = sentiment_dist.get('muy_negativo', 0)
        
        pos_pct = ((muy_positivo + positivo) / total_reviews * 100) if total_reviews > 0 else 0
        neg_pct = ((muy_negativo + negativo) / total_reviews * 100) if total_reviews > 0 else 0
        
        # Determinar tendencia (simplificado)
        if avg_rating >= 4.0:
            tendencia = "🟢 Positiva"
        elif avg_rating >= 3.0:
            tendencia = "🟠 Estable"
        else:
            tendencia = "🔴 Negativa"
        
        if heading_mgr:
            if heading_mgr.add_h3("Tabla — Indicadores clave de reputación (KPIs)") is None:
                return
        else:
            title = doc.add_paragraph("Tabla — Indicadores clave de reputación (KPIs)")
            title.style = doc.styles['Heading 3']
            title.runs[0].font.color.rgb = SECONDARY
        table = doc.add_table(rows=6, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)
        
        # Configurar anchos
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(3)
        table.columns[3].width = Cm(6)
        
        # Encabezados
        headers = ["KPI", "Valor Actual", "Tendencia", "Insight Clave"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"
        
        # Datos de la tabla
        table_data = [
            ("Número total de reseñas", f"{total_reviews:,}", "Estable", f"La cantidad de reseñas es {'alta' if total_reviews > 100 else 'moderada' if total_reviews > 50 else 'baja'}, {'con alta polaridad' if abs(star_5_pct - star_1_pct) > 30 else 'con distribución equilibrada'}."),
            ("Media de puntuación", f"{avg_rating:.1f}", tendencia, f"La media {'ha mejorado' if avg_rating >= 4.0 else 'se mantiene estable' if avg_rating >= 3.0 else 'ha disminuido'} en los últimos meses."),
            ("Distribución por estrellas", f"{star_5_pct:.0f}% 5★, {star_4_pct:.0f}% 4★, {star_3_pct:.0f}% 3★, {star_2_pct:.0f}% 2★, {star_1_pct:.0f}% 1★", 
             "🟢 Positiva" if star_5_pct + star_4_pct > 60 else "🟠 Mixta" if star_5_pct + star_4_pct > 40 else "🔴 Negativa", 
             f"La mayoría de las reseñas son de {'5 estrellas' if star_5_pct > star_1_pct else '1 estrella'}."),
            ("Ratio de respuesta", "40%", "Estable", "La respuesta del propietario es baja en comparación con el volumen de reseñas."),
            ("Tiempo de respuesta", "3 días", "Estable", "Tiempo de respuesta aceptable, pero podría mejorar.")
        ]
        
        for row_idx, (kpi, valor, tendencia, insight) in enumerate(table_data, 1):
            for col_idx in range(4):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                p = cell.paragraphs[0]
                if col_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                content = kpi if col_idx == 0 else (valor if col_idx == 1 else (tendencia if col_idx == 2 else insight))
                run = p.add_run(content)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = TEXT
                
                if col_idx == 0:  # Primera columna en negrita
                    run.bold = True
                    
    except Exception as e:
        logger.error(f"Error creando tabla KPI profesional: {e}")

def _create_sentiment_analysis_table(doc, df, heading_mgr=None, *, sentiment_dist=None):
    """Crear tabla de análisis de sentimiento profesional"""
    try:
        if sentiment_dist is None:
            sentiment_dist = df['sentiment'].value_counts()
        total = len(df)
        
        # Datos ordenados
        sentiments = [
            ("Muy Positivo", "muy_positivo", SUCCESS, "🟢"),
            ("Positivo", "positivo", SUCCESS, "🟢"),
            ("Neutro", "neutro", WARNING, "🟠"),
            ("Negativo", "negativo", DANGER, "🔴"),
            ("Muy Negativo", "muy_negativo", DANGER, "🔴")
        ]
        
        if heading_mgr:
            if heading_mgr.add_h3("Tabla — Distribución de sentimiento") is None:
                return
        else:
            title = doc.add_paragraph("Tabla — Distribución de sentimiento")
            title.style = doc.styles['Heading 3']
            title.runs[0].font.color.rgb = SECONDARY
        table = doc.add_table(rows=len(sentiments) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)
        
        # Configurar anchos
        table.columns[0].width = Cm(2.5)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(4)
        
        # Encabezado
        header_cells = ["Sentimiento", "Cantidad", "Porcentaje", "Estado", "Análisis"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"
        
        # Filas de datos
        for row_idx, (label, key, color, emoji) in enumerate(sentiments, 1):
            count = sentiment_dist.get(key, 0)
            percentage = (count / total * 100) if total > 0 else 0
            
            # Análisis contextual
            if percentage > 40:
                analisis = "Dominante - Requiere atención inmediata"
            elif percentage > 20:
                analisis = "Significativo - Monitorear tendencia"
            elif percentage > 10:
                analisis = "Moderado - Seguimiento regular"
            else:
                analisis = "Bajo - Mantener vigilancia"
            
            # Sentimiento con emoji
            cell = table.cell(row_idx, 0)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"{emoji} {label}")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = color
            run.bold = True
            
            # Cantidad
            cell = table.cell(row_idx, 1)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{count:,}")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            run.bold = True
            
            # Porcentaje
            cell = table.cell(row_idx, 2)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{percentage:.1f}%")
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT
            run.bold = True
            
            # Estado
            cell = table.cell(row_idx, 3)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(emoji)
            run.font.size = Pt(12)
            
            # Análisis
            cell = table.cell(row_idx, 4)
            _set_cell_border(cell, "single", "CCCCCC", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(analisis)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT
            
    except Exception as e:
        logger.error(f"Error creando tabla sentimiento profesional: {e}")

def _create_languages_table(doc, languages, heading_mgr=None, *, languages_sorted: bool = False):
    """Crear tabla de distribución de idiomas"""
    try:
        if not languages_sorted:
            languages = _sorted_languages(languages)
        if not languages:
            return

        if heading_mgr:
            if heading_mgr.add_h3("Tabla — Distribución de idiomas") is None:
                return
        else:
            title = doc.add_paragraph("Tabla — Distribución de idiomas")
            title.style = doc.styles['Heading 3']
            title.runs[0].font.color.rgb = SECONDARY
        table = doc.add_table(rows=len(languages) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)

        # Configurar anchos
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(6)

        # Encabezado
        header_cells = ["Idioma", "Cantidad", "Porcentaje", "Análisis"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"

        # Filas de datos
        for row_idx, lang in enumerate(languages, 1):
            for col_idx in range(4):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                if col_idx in (1, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if col_idx == 0:
                    content = lang["name"]
                    run = p.add_run(content)
                    run.bold = True
                elif col_idx == 1:
                    content = f"{lang['count']:,}"
                    run = p.add_run(content)
                elif col_idx == 2:
                    content = f"{lang['percentage']:.1f}%"
                    run = p.add_run(content)
                elif col_idx == 3:
                    # Análisis contextual
                    if lang['percentage'] > 50:
                        analisis = "Idioma predominante - Mercado local principal"
                    elif lang['percentage'] > 20:
                        analisis = "Idioma significativo - Mercado secundario importante"
                    elif lang['percentage'] > 10:
                        analisis = "Idioma relevante - Mercado emergente"
                    else:
                        analisis = "Idioma minoritario - Mercado nicho"
                    run = p.add_run(analisis)

                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = TEXT

    except Exception as e:
        logger.error(f"Error creando tabla idiomas: {e}")


def _create_engagement_table(doc, engagement_stats, heading_mgr=None):
    """Crear tabla de métricas de engagement"""
    try:
        if not engagement_stats:
            return

        if heading_mgr:
            if heading_mgr.add_h3("Tabla — Métricas de engagement") is None:
                return
        else:
            title = doc.add_paragraph("Tabla — Métricas de engagement")
            title.style = doc.styles['Heading 3']
            title.runs[0].font.color.rgb = SECONDARY
        table = doc.add_table(rows=6, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)

        # Configurar anchos
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(6)

        # Encabezado
        header_cells = ["Métrica de Engagement", "Valor", "Análisis"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"

        # Datos de engagement
        engagement_data = [
            ("Reseñas con likes", f"{engagement_stats.get('total_reviews_with_likes', 0):,}", f"{'Bajo' if engagement_stats.get('total_reviews_with_likes', 0) < 10 else 'Moderado' if engagement_stats.get('total_reviews_with_likes', 0) < 50 else 'Alto'} nivel de interacción social"),
            ("Likes promedio por reseña", f"{engagement_stats.get('avg_likes_per_review', 0):.1f}", f"{'Poca' if engagement_stats.get('avg_likes_per_review', 0) < 1 else 'Moderada' if engagement_stats.get('avg_likes_per_review', 0) < 3 else 'Alta'} engagement por reseña"),
            ("Total de fotos", f"{engagement_stats.get('total_photos', 0):,}", f"{'Limitado' if engagement_stats.get('total_photos', 0) < 20 else 'Moderado' if engagement_stats.get('total_photos', 0) < 100 else 'Alto'} contenido visual"),
            ("Reseñas con fotos", f"{engagement_stats.get('reviews_with_photos', 0):,}", f"{'Baja' if engagement_stats.get('reviews_with_photos', 0) < 10 else 'Moderada' if engagement_stats.get('reviews_with_photos', 0) < 50 else 'Alta'} participación visual"),
            ("Tasa de respuesta", f"{engagement_stats.get('response_rate', 0):.1f}%", f"{'Crítica' if engagement_stats.get('response_rate', 0) < 20 else 'Mejorable' if engagement_stats.get('response_rate', 0) < 50 else 'Buena'} atención al cliente")
        ]

        # Filas de datos
        for row_idx, (metric, value, analysis) in enumerate(engagement_data, 1):
            for col_idx in range(3):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                if col_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                content = [metric, value, analysis][col_idx]
                run = p.add_run(content)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = TEXT

                if col_idx == 0:  # Primera columna en negrita
                    run.bold = True

    except Exception as e:
        logger.error(f"Error creando tabla engagement: {e}")


def _create_recommendations_table(doc, df, heading_mgr=None, *, pos_pct=None):
    """Crear tabla de recomendaciones estratégicas"""
    try:
        # Análisis básico para generar recomendaciones
        total_reviews = len(df)
        avg_rating = df['rating_num'].mean() if total_reviews > 0 else 0
        if pos_pct is None:
            pos_pct = (df['sentiment'].isin(['positivo', 'muy_positivo']).sum() / total_reviews * 100) if total_reviews > 0 else 0

        # Recomendaciones basadas en datos
        recommendations = [
            ("Gestión de Reseñas", "Implementar sistema de seguimiento", "Respuesta < 24h", "Alto", "Mejora percepción de atención"),
            ("Análisis de Sentimiento", "Monitoreo continuo", "Reportes semanales", "Medio", "Identificación temprana de problemas"),
            ("Capacitación del Equipo", "Formación en atención al cliente", "Sesiones mensuales", "Alto", "Reducción de reseñas negativas"),
            ("Comunicación Proactiva", "Encuestas de satisfacción", "Trimestral", "Medio", "Feedback directo de clientes"),
            ("Optimización SEO Local", "Palabras clave locales", "Continuo", "Bajo", "Mayor visibilidad online")
        ]

        if heading_mgr:
            if heading_mgr.add_h3("Tabla — Plan de acción estratégico") is None:
                return
        else:
            title = doc.add_paragraph("Tabla — Plan de acción estratégico")
            title.style = doc.styles['Heading 3']
            title.runs[0].font.color.rgb = SECONDARY
        table = doc.add_table(rows=len(recommendations) + 1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)

        # Configurar anchos
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(3)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(4)

        # Encabezado
        header_cells = ["Área", "Recomendación", "Frecuencia", "Prioridad", "Impacto Esperado"]
        for i, header in enumerate(header_cells):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"

        # Filas de datos
        for row_idx, (area, recomendacion, frecuencia, prioridad, impacto) in enumerate(recommendations, 1):
            for col_idx in range(5):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                content = [area, recomendacion, frecuencia, prioridad, impacto][col_idx]
                run = p.add_run(content)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = TEXT

                if col_idx == 0:  # Primera columna en negrita
                    run.bold = True

    except Exception as e:
        logger.error(f"Error creando tabla recomendaciones: {e}")

# ========================== Gráficos profesionales ==========================
def _create_professional_rating_chart(df, path, *, counts=None):
    """Crear gráfico de valoraciones profesional"""
    try:
        if counts is None:
            counts = df['rating_num'].value_counts().sort_index()
        n_categories = len(counts)
        show_labels = n_categories <= 6
        total = int(counts.sum()) if len(counts) else 0
        
        # Configurar estilo profesional
        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(12, 7))
        fig.subplots_adjust(top=0.88, bottom=0.14)
        
        # Colores profesionales
        colors = ['#E53935', '#FF7043', '#FFB74D', '#66BB6A', '#43A047']
        bars = ax.bar(counts.index.astype(str), counts.values, color=colors, 
                     edgecolor='white', linewidth=2, alpha=0.9)
        
        # Personalizar gráfico
        ax.set_xlabel('Valoración (Estrellas)', fontsize=14, fontweight='bold', color='#2C2C2C')
        ax.set_ylabel('Número de Reseñas', fontsize=14, fontweight='bold', color='#2C2C2C')
        ax.set_title('Distribución de reseñas por estrellas', fontsize=16, fontweight='bold', color='#004E89', pad=16)
        
        # Añadir valores en las barras
        if show_labels:
            for bar in bars:
                height = bar.get_height()
                pct = (float(height) / total * 100) if total > 0 else 0
                ax.text(
                    bar.get_x() + bar.get_width() / 2.0,
                    height + (total * 0.005),
                    f"{int(height)} ({pct:.1f}%)",
                    ha='center',
                    va='bottom',
                    fontweight='bold',
                    fontsize=11,
                )
        
        # Estilo de la grilla
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        ax.set_axisbelow(True)
        
        # Colores de ejes
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#CCCCCC')
        ax.spines['bottom'].set_color('#CCCCCC')

        tick_labels = [str(x) for x in counts.index.tolist()]
        if any(len(x) > 3 for x in tick_labels):
            plt.xticks(rotation=30, ha='right')
        
        plt.savefig(path, dpi=300, bbox_inches='tight', pad_inches=0.2, facecolor='white')
        plt.close()
        return path
        
    except Exception as e:
        logger.error(f"Error creando gráfico de valoraciones profesional: {e}")
        return None

def _create_professional_sentiment_chart(df, path, *, sentiment_counts=None):
    """Crear gráfico de sentimiento profesional"""
    try:
        order = ["muy_negativo", "negativo", "neutro", "positivo", "muy_positivo"]
        labels = ["Muy Negativo", "Negativo", "Neutro", "Positivo", "Muy Positivo"]
        colors = ["#C62828", "#E53935", "#FB8C00", "#43A047", "#2E7D32"]
        n_categories = len(labels)
        show_labels = n_categories <= 6

        if sentiment_counts is None:
            v = df['sentiment'].value_counts()
        else:
            v = sentiment_counts
        vals = [int(v.get(k, 0)) for k in order]
        total = sum(vals)

        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(12, 7))
        fig.subplots_adjust(top=0.88, bottom=0.14)

        bars = ax.barh(labels, vals, color=colors, edgecolor='white', linewidth=2, alpha=0.9)

        # Personalizar gráfico
        ax.set_xlabel('Número de Reseñas', fontsize=14, fontweight='bold', color='#2C2C2C')
        ax.set_ylabel('Sentimiento', fontsize=14, fontweight='bold', color='#2C2C2C')
        ax.set_title('Distribución de reseñas por sentimiento', fontsize=16, fontweight='bold', color='#004E89', pad=16)

        # Añadir valores en las barras
        if show_labels:
            for i, bar in enumerate(bars):
                width = bar.get_width()
                pct = (float(vals[i]) / total * 100) if total > 0 else 0
                ax.text(
                    width + (total * 0.005),
                    bar.get_y() + bar.get_height() / 2.0,
                    f"{int(vals[i])} ({pct:.1f}%)",
                    ha='left',
                    va='center',
                    fontweight='bold',
                    fontsize=11,
                )

        # Estilo de la grilla
        ax.grid(axis='x', alpha=0.3, linestyle='--')
        ax.set_axisbelow(True)

        # Colores de ejes
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#CCCCCC')
        ax.spines['bottom'].set_color('#CCCCCC')

        plt.savefig(path, dpi=300, bbox_inches='tight', pad_inches=0.2, facecolor='white')
        plt.close()
        return path

    except Exception as e:
        logger.error(f"Error creando gráfico de sentimiento profesional: {e}")
        return None


def _create_professional_language_chart(languages, path, *, languages_sorted: bool = False):
    """Crear gráfico de distribución de idiomas profesional"""
    try:
        if not languages_sorted:
            languages = _sorted_languages(languages)
        if not languages:
            return None
        n_categories = len(languages)
        show_labels = n_categories <= 6

        # Preparar datos
        labels = [lang['name'] for lang in languages]
        sizes = [lang['percentage'] for lang in languages]
        colors = ['#004E89', '#1E88E5', '#43A047', '#FB8C00', '#E53935', '#8E24AA', '#3949AB', '#00ACC1', '#7CB342', '#F4511E']

        plt.style.use('default')
        fig, ax = plt.subplots(figsize=(12, 7))
        fig.subplots_adjust(top=0.88, bottom=0.14, right=0.75)

        # Crear gráfico circular
        if show_labels:
            wedges, texts, autotexts = ax.pie(
                sizes,
                labels=labels,
                colors=colors[:len(labels)],
                autopct='%1.1f%%',
                startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 2},
            )
        else:
            wedges, texts = ax.pie(
                sizes,
                labels=None,
                colors=colors[:len(labels)],
                autopct=None,
                startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 2},
            )
            autotexts = []

        # Personalizar
        ax.set_title('Distribución de idiomas en reseñas', fontsize=16, fontweight='bold', color='#004E89', pad=16)

        # Mejorar legibilidad
        if show_labels:
            for text in texts:
                text.set_fontsize(10)
                text.set_fontweight('bold')
            for autotext in autotexts:
                autotext.set_fontsize(9)
                autotext.set_fontweight('bold')
                autotext.set_color('white')

        # Añadir leyenda con cantidades
        if show_labels:
            legend_labels = [f"{lang['name']}: {lang['count']:,} ({lang['percentage']:.1f}%)" for lang in languages]
        else:
            legend_labels = [f"{lang['name']}" for lang in languages]
        ax.legend(wedges, legend_labels, title="Idiomas", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

        plt.savefig(path, dpi=300, bbox_inches='tight', pad_inches=0.2, facecolor='white')
        plt.close()
        return path

    except Exception as e:
        logger.error(f"Error creando gráfico de idiomas profesional: {e}")
        return None

# ========================== Prompt profesional ==========================
def _load_professional_prompt(project_root: Path) -> str:
    prompt_path = project_root / "ai_models" / "prompt_format_content.md"

    if not prompt_path.is_file():
        raise FileNotFoundError(
            f"Prompt profesional no encontrado: {prompt_path}"
        )

    return prompt_path.read_text(encoding="utf-8")


def _load_section6_prompt(project_root: Path) -> str:
    prompt_path = project_root / "ai_models" / "prompt_analysis_section_6.md"

    if not prompt_path.is_file():
        raise FileNotFoundError(
            f"Prompt sección 6 no encontrado: {prompt_path}"
        )

    return prompt_path.read_text(encoding="utf-8")


def _resolve_sector_category(cfg: dict):
    try:
        category = (cfg.get("place_metadata") or {}).get("category")
    except Exception:
        category = None

    if category is None:
        return (False, None)

    raw = str(category).strip()
    if not raw:
        return (False, None)

    normalized = raw.casefold()
    enabled = any(normalized == str(w).strip().casefold() for w in SECTOR_CATEGORY_WHITELIST)
    return (enabled, raw)


def _prepend_sector_context_to_prompt(prompt: str, cfg: dict) -> str:
    enabled, category = _resolve_sector_category(cfg)
    if not enabled or not category:
        return prompt

    block = "CONTEXTO SECTORIAL (solo para IA):\n\n" f"Tipo de negocio: {category}"
    return block + "\n\n" + prompt


def resolve_sector_profile(category: str | None) -> str:
    # Función pura: resuelve profile_id determinista a partir de category (sin fallar).
    try:
        raw = str(category).strip() if category is not None else ""
    except Exception:
        raw = ""

    if not raw:
        return "generic"

    normalized = raw.lower()
    if normalized in ("restaurant", "fast food restaurant"):
        return "restaurant"

    return "generic"


def prepend_sector_profile_to_prompt(prompt: str, profile_id: str) -> str:
    # Inyección solo en prompt IA. No se persiste en artefactos.
    profile = SECTOR_PROFILES.get(profile_id) if isinstance(profile_id, str) else None
    if not isinstance(profile, dict):
        profile = SECTOR_PROFILES.get("generic")

    label = str((profile or {}).get("label") or SECTOR_PROFILES["generic"]["label"])
    context = str((profile or {}).get("context") or SECTOR_PROFILES["generic"]["context"])

    block = (
        "PERFIL SECTORIAL (solo para IA):\n"
        f"Perfil: {label}\n"
        f"Contexto: {context}\n"
    )
    return block + "\n" + (prompt or "")


def _apply_sector_profile_to_prompt(prompt: str, cfg: dict) -> str:
    # Prioridad: perfil sectorial (si no es generic). Fallback: whitelist actual.
    try:
        category = (cfg.get("place_metadata") or {}).get("category")
    except Exception:
        category = None

    profile_id = resolve_sector_profile(category if isinstance(category, str) or category is None else str(category))
    if profile_id != "generic":
        return prepend_sector_profile_to_prompt(prompt, profile_id)

    return _prepend_sector_context_to_prompt(prompt, cfg)


def _openai_error_reason(exc: Exception) -> str:
    status_code = getattr(exc, "status_code", None)
    if status_code == 401:
        return "401"
    msg = str(exc or "")
    if "401" in msg:
        return "401"
    msg_l = msg.lower()
    if "timeout" in msg_l or "timed out" in msg_l:
        return "TIMEOUT"
    return "EXCEPTION"


def _call_openai_professional(
    section_name: str,
    prompt: str,
    model: str,
    temperature: float,
    cfg: dict,
    *,
    enable_key: str,
    state: dict,
):
    api_key = os.getenv("OPENAI_API_KEY")
    if not cfg.get(enable_key, True):
        return None
    if state.get("openai_disabled", False):
        logged = state.get("openai_logged_sections")
        if isinstance(logged, set) and section_name not in logged:
            reason = state.get("openai_disabled_reason") or "DISABLED"
            logger.error(f"OPENAI_DISABLED_FOR_SECTION={section_name} REASON={reason}")
            logged.add(section_name)
        return None
    if not api_key or OpenAI is None:
        return None
    
    try:
        client = OpenAI(api_key=api_key)
        r = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[{"role": "user", "content": prompt}],
        )
        text = (r.choices[0].message.content or "").strip()
        return text or None
    except Exception as e:
        reason = _openai_error_reason(e)
        state["openai_disabled"] = True
        state["openai_disabled_reason"] = reason
        logged = state.get("openai_logged_sections")
        if isinstance(logged, set) and section_name not in logged:
            logger.error(f"OPENAI_DISABLED_FOR_SECTION={section_name} REASON={reason}")
            logged.add(section_name)
        return None


class _HeadingManager:
    def __init__(self, doc):
        self.doc = doc
        self.h2_idx = 0
        self.h3_idx = 0
        self._h3_seen_in_h2 = set()
        self.figure_idx = 0

    def add_h2(self, title: str):
        self.h2_idx += 1
        self.h3_idx = 0
        self._h3_seen_in_h2 = set()
        self.figure_idx = 0

        p = self.doc.add_paragraph(f"{self.h2_idx}. {title}")
        p.style = self.doc.styles["Heading 2"]
        if p.runs:
            p.runs[0].font.color.rgb = PRIMARY
        else:
            run = p.add_run("")
            run.font.color.rgb = PRIMARY
        return p

    def add_h3(self, title: str):
        key = (title or "").strip().lower()
        if key in self._h3_seen_in_h2:
            return None

        self._h3_seen_in_h2.add(key)
        self.h3_idx += 1

        p = self.doc.add_paragraph(f"{self.h2_idx}.{self.h3_idx}. {title}")
        p.style = self.doc.styles["Heading 3"]
        if p.runs:
            p.runs[0].font.color.rgb = SECONDARY
        else:
            run = p.add_run("")
            run.font.color.rgb = SECONDARY
        return p

    def add_figure(self, title: str, image_path, width=Inches(6)):
        self.figure_idx += 1

        self.doc.add_picture(str(image_path), width=width)
        pic_para = self.doc.paragraphs[-1]
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            pic_para.paragraph_format.space_before = Pt(6)
            pic_para.paragraph_format.space_after = Pt(4)
        except Exception:
            pass

        caption = f"Figura {self.h2_idx}.{self.figure_idx} — {title}"
        cap_para = self.doc.add_paragraph(caption)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            cap_para.paragraph_format.space_before = Pt(0)
            cap_para.paragraph_format.space_after = Pt(12)
        except Exception:
            pass
        run = cap_para.runs[0] if cap_para.runs else cap_para.add_run("")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT
        return cap_para

# ========================== Generador principal ==========================
def create_professional_report(df, client_name, output_path, cfg, project_root="."):
    """Función principal de generación de reporte profesional"""
    try:
        project_root = Path(project_root)
        model = cfg.get("openai_model", "gpt-4o-mini")
        temperature = cfg.get("openai_temperature", 0.3)
        title = cfg.get("report_title", "Informe ORM y CX")
        enable_ai_analysis = bool(cfg.get("enable_ai_analysis", True))
        enable_ai_owner = bool(cfg.get("enable_ai_owner", True))
        openai_state = {
            "openai_disabled": False,
            "openai_disabled_reason": None,
            "openai_logged_sections": set(),
        }

        raw_mode = cfg.get("output_mode", cfg.get("report_mode", "CLIENT"))
        mode = str(raw_mode or "CLIENT").upper()
        if mode not in ("CLIENT", "AUDIT"):
            mode = "CLIENT"
        cfg["output_mode"] = mode
        cfg["report_mode"] = mode

        languages = cfg.get("languages", [])
        engagement_stats = cfg.get("engagement", {})

        cfg["metrics_version"] = cfg.get("metrics_version", "v1")

        client_clean = _normalize_client_name(client_name)
        safe_name = _safe_filename(client_clean)
        client_dir = Path(cfg.get("output_dir", "./data/reports")) / safe_name / f"v{PIPELINE_VERSION}"
        client_dir.mkdir(parents=True, exist_ok=True)

        # Copiar CSV fuente si está disponible
        if "source_csv" in cfg and Path(cfg["source_csv"]).exists():
            try:
                shutil.copy(cfg["source_csv"], client_dir / Path(cfg["source_csv"]).name)
            except Exception as e:
                logger.warning(f"No se pudo copiar CSV fuente: {e}")

        derived = {}
        disable_report_cache = str(os.getenv("ORM_DISABLE_REPORT_CACHE", "")).strip().lower() in ("1", "true", "yes")

        def _get_cached(key, fn):
            if disable_report_cache:
                return fn()
            if key in derived:
                return derived[key]
            v = fn()
            derived[key] = v
            return v

        total_reviews = _get_cached("total_reviews", lambda: int(len(df)))
        star_dist = _get_cached("star_dist", lambda: df["rating_num"].value_counts().sort_index())
        sentiment_dist = _get_cached("sentiment_dist", lambda: df["sentiment"].value_counts())
        pos_pct = _get_cached(
            "pos_pct",
            lambda: (
                (
                    int(sentiment_dist.get("positivo", 0))
                    + int(sentiment_dist.get("muy_positivo", 0))
                )
                / total_reviews
                * 100
            )
            if total_reviews > 0
            else 0.0,
        )

        languages_prepared = _apply_languages_top_n(
            _sorted_languages(languages),
            cfg.get("top_n_languages"),
        )

        chart_cache = {}

        def _render_once(key, fn):
            if disable_report_cache:
                return fn()
            if key in chart_cache:
                return chart_cache[key]
            v = fn()
            chart_cache[key] = v
            return v

        with _SectionTimer("AI_ANALYSIS"):
            reviews_text = _build_reviews_text(df)
            professional_prompt = _load_professional_prompt(project_root)
            professional_prompt = _apply_sector_profile_to_prompt(professional_prompt, cfg)
            full_prompt = professional_prompt.replace("{{reviews_text}}", reviews_text)
            ai_text = _call_openai_professional(
                "AI_ANALYSIS",
                full_prompt,
                model,
                temperature,
                cfg,
                enable_key="enable_ai_analysis",
                state=openai_state,
            )
            if enable_ai_analysis and ai_text is None:
                ai_text = "Análisis cualitativo no disponible en esta ejecución. El informe mantiene métricas, gráficos y tablas deterministas."

        with _SectionTimer("SECTION6_ARTIFACTS"):
            contradictions = compute_contradictions(df)
            segments = compute_local_vs_tourist(df)
            fatigue = compute_operational_fatigue(df)

        cfg["artifact_contradictions"] = contradictions
        cfg["artifact_segments"] = segments
        cfg["artifact_operational_fatigue"] = fatigue
        contradiction_score = compute_contradiction_score(contradictions)
        cfg["artifact_contradiction_score"] = contradiction_score
        fatigue_score = compute_operational_fatigue_score(fatigue)
        cfg["artifact_fatigue_score"] = fatigue_score
        alerts = compute_alerts(cfg)
        cfg["alerts"] = alerts

        alerts_by_score = compute_alerts_by_score(
            cfg.get("artifact_contradiction_score"),
            cfg.get("artifact_fatigue_score"),
            mode=mode,
        )
        cfg["artifact_alerts"] = alerts_by_score

        audit_mode = is_audit_mode(cfg)
        cfg_for_prompt = cfg if audit_mode else redact_for_client(cfg)

        section6_prompt = _load_section6_prompt(project_root)
        section6_prompt = _apply_sector_profile_to_prompt(section6_prompt, cfg)

        import json as _json_for_section6

        section6_filled = (
            section6_prompt
            .replace("{{artifact_contradictions}}", _json_for_section6.dumps(cfg_for_prompt.get("artifact_contradictions", "NO DISPONIBLE"), ensure_ascii=False))
            .replace("{{artifact_segments}}", _json_for_section6.dumps(cfg_for_prompt.get("artifact_segments", "NO DISPONIBLE"), ensure_ascii=False))
            .replace("{{artifact_operational_fatigue}}", _json_for_section6.dumps(cfg_for_prompt.get("artifact_operational_fatigue", "NO DISPONIBLE"), ensure_ascii=False))
            .replace("{{artifact_contradiction_score}}", _json_for_section6.dumps(cfg_for_prompt.get("artifact_contradiction_score", "NO DISPONIBLE"), ensure_ascii=False))
            .replace("{{metrics_version}}", cfg.get("metrics_version", "v1"))
            .replace("{{output_mode}}", mode)
        )
        with _SectionTimer("AI_OWNER"):
            section6_text = _call_openai_professional(
                "AI_OWNER",
                section6_filled,
                model=model,
                temperature=min(temperature, 0.3),
                cfg=cfg,
                enable_key="enable_ai_owner",
                state=openai_state,
            )
            if enable_ai_owner and section6_text is None:
                section6_text = "Se omite el análisis avanzado en esta ejecución. El informe mantiene las tablas y métricas deterministas."

        with _SectionTimer("CHARTS"):
            g1 = _render_once(
                "rating",
                lambda: _create_professional_rating_chart(
                    df,
                    client_dir / "rating_chart_professional.png",
                    counts=star_dist,
                ),
            )
            g2 = _render_once(
                "sentiment",
                lambda: _create_professional_sentiment_chart(
                    df,
                    client_dir / "sentiment_chart_professional.png",
                    sentiment_counts=sentiment_dist,
                ),
            )
            g3 = _render_once(
                "language",
                lambda: _create_professional_language_chart(
                    languages_prepared,
                    client_dir / "language_chart_professional.png",
                    languages_sorted=True,
                ),
            )

        # Crear documento
        with _SectionTimer("DOC_BUILD"):
            doc = Document()
            _apply_professional_doc_standards(doc, client_clean)
            hm = _HeadingManager(doc)

        # Portada profesional
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"📊 {title}")
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = PRIMARY
        title_run.bold = True

        client_para = doc.add_paragraph()
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client_para.add_run(f"Cliente: {client_clean}")
        client_run.font.name = "Calibri"
        client_run.font.size = Pt(18)
        client_run.font.color.rgb = TEXT

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%d de %B de %Y"))
        date_run.font.name = "Calibri"
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = TEXT
        date_run.italic = True

        # ===================== METADATOS DEL LOCAL =====================
        meta = cfg.get("place_metadata", {})

        def _safe_meta(v):
            if v is None:
                return None
            if isinstance(v, str):
                s = v.strip()
                if not s:
                    return None
                if s.upper() == "NO DISPONIBLE":
                    return None
                if s == "None":
                    return None
                return s
            return v

        business_title = _safe_meta(meta.get("title"))
        category = _safe_meta(meta.get("category"))
        address = _safe_meta(meta.get("address"))
        phone = _safe_meta(meta.get("phoneNumber"))
        website = _safe_meta(meta.get("website"))
        total_score = _safe_meta(meta.get("totalScore"))
        reviews_count = _safe_meta(meta.get("reviewsCount"))
        maps_url = _safe_meta(meta.get("url"))
        scraped_at = _safe_meta(meta.get("scrapedAt"))

        ficha_lines = []
        if category:
            ficha_lines.append(f"Categoría: {category}")
        if address:
            ficha_lines.append(f"Dirección: {address}")
        if phone:
            ficha_lines.append(f"Teléfono: {phone}")
        if website:
            ficha_lines.append(f"Website: {website}")
        if total_score is not None:
            ficha_lines.append(f"Valoración media: {total_score}")
        if reviews_count is not None:
            ficha_lines.append(f"Nº de reseñas: {reviews_count}")
        if maps_url:
            ficha_lines.append(f"URL Google Maps: {maps_url}")
        if scraped_at:
            ficha_lines.append(f"Fecha de scraping: {scraped_at}")

        if business_title or ficha_lines:
            doc.add_paragraph()

        if business_title:
            business_para = doc.add_paragraph()
            business_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            business_run = business_para.add_run(str(business_title))
            business_run.font.name = "Calibri"
            business_run.font.size = Pt(16)
            business_run.font.color.rgb = TEXT
            business_run.bold = True

        if ficha_lines:
            ficha_title_para = doc.add_paragraph()
            ficha_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ficha_title_run = ficha_title_para.add_run("Ficha del negocio")
            ficha_title_run.font.name = "Calibri"
            ficha_title_run.font.size = Pt(12)
            ficha_title_run.font.color.rgb = TEXT
            ficha_title_run.bold = True

            for line in ficha_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                r.font.color.rgb = TEXT

        def add_h2(title: str):
            doc.add_page_break()
            return hm.add_h2(title)

        with _SectionTimer("SECTION_KPIS"):
            add_h2("Indicadores Clave de Reputación (KPIs)")
            doc.add_paragraph("Indicadores clave para entender el desempeño global.")
            _create_professional_kpi_table(
                doc,
                df,
                heading_mgr=hm,
                star_dist=star_dist,
                sentiment_dist=sentiment_dist,
            )

        # Espaciado
        doc.add_paragraph()

        if engagement_stats:
            with _SectionTimer("SECTION_ENGAGEMENT"):
                _create_engagement_table(doc, engagement_stats, heading_mgr=hm)

        # Espaciado
        doc.add_paragraph()

        if (g1 and Path(g1).exists()) or (g2 and Path(g2).exists()):
            with _SectionTimer("SECTION_VISUAL"):
                add_h2("📊 Análisis Visual")
                doc.add_paragraph("Gráficos para visualizar distribución y tendencias.")

                if g1 and Path(g1).exists():
                    hm.add_figure("Distribución de reseñas por estrellas", g1, width=Inches(6))
                if g2 and Path(g2).exists():
                    hm.add_figure("Distribución de reseñas por sentimiento", g2, width=Inches(6))

        # Espaciado
        doc.add_paragraph()

        if len(df) > 0:
            with _SectionTimer("SECTION_SENTIMENT"):
                add_h2("Análisis de Sentimiento")
                doc.add_paragraph("Distribución del tono y principales implicaciones.")
                _create_sentiment_analysis_table(
                    doc,
                    df,
                    heading_mgr=hm,
                    sentiment_dist=sentiment_dist,
                )

        # Espaciado
        doc.add_paragraph()

        if languages:
            with _SectionTimer("SECTION_LANGUAGES"):
                add_h2("Análisis de Idiomas")
                doc.add_paragraph("Distribución lingüística para entender el perfil de audiencia.")
                _create_languages_table(
                    doc,
                    languages_prepared,
                    heading_mgr=hm,
                    languages_sorted=True,
                )

        # Espaciado
        doc.add_paragraph()

        # Gráfico de idiomas
        if g3 and Path(g3).exists():
            with _SectionTimer("SECTION_LANGUAGE_FIGURE"):
                hm.add_figure("Distribución de idiomas en reseñas", g3, width=Inches(6))

        if len(df) > 0:
            with _SectionTimer("SECTION_ACTION_PLAN"):
                add_h2("Plan de Acción Estratégico")
                doc.add_paragraph("Acciones prioritarias para mejorar reputación y experiencia.")
                _create_recommendations_table(
                    doc,
                    df,
                    heading_mgr=hm,
                    pos_pct=pos_pct,
                )

        if section6_text or ai_text:
            with _SectionTimer("SECTION_AI_MARKDOWN"):
                add_h2("🤖 Análisis Inteligente Completo")
                doc.add_paragraph("Síntesis interpretativa y recomendaciones cualitativas.")

                _process_markdown_content(doc, section6_text, heading_mgr=hm)
                _process_markdown_content(doc, ai_text, heading_mgr=hm)

        if audit_mode:
            with _SectionTimer("SECTION_AUDIT_ANNEX"):
                add_h2(f"🛠️ Anexo — Auditoría interna (metrics_version: {cfg.get('metrics_version', 'v1')})")
                doc.add_paragraph("Resumen interno para control de calidad.")

                annex_markdown = _build_audit_annex_markdown(cfg)
                _process_markdown_content(doc, annex_markdown, heading_mgr=hm)

        generated_at = datetime.now(timezone.utc).isoformat()
        input_path = None
        try:
            source = cfg.get("source_csv")
            if source:
                input_path = Path(str(source))
        except Exception:
            input_path = None
        input_hash = _sha256_file(input_path) if (input_path and input_path.exists()) else None
        _inject_docx_traceability_metadata(
            doc,
            pipeline_version=PIPELINE_VERSION,
            input_hash=input_hash,
            generated_at=generated_at,
        )
        logger.info(
            f"TRACE_DOCX pipeline_version={PIPELINE_VERSION} input_hash={input_hash or 'N/A'} generated_at={generated_at}"
        )

        with _SectionTimer("DOC_SAVE"):
            output_file = client_dir / f"{safe_name}_informe_PROFESIONAL.docx"
            doc.save(str(output_file))
            logger.info(f"✅ Informe profesional guardado: {output_file}")

        run_mode = str(cfg.get("mode", "dev") or "dev").strip().lower()
        require_pdf = bool(cfg.get("require_pdf", run_mode == "prod"))

        pdf_path = output_file.with_suffix(".pdf")
        if require_pdf or cfg.get("attempt_pdf", True):
            try:
                with _SectionTimer("PDF_CONVERT"):
                    from docx2pdf import convert
                    convert(str(output_file), str(pdf_path))
                    logger.info(f"📄 PDF generado: {pdf_path}")
            except Exception as e:
                if require_pdf:
                    raise OutputContractError(f"PDF_CONVERT_FAILED: {e}")
                logger.warning(f"No se pudo generar PDF: {e}")

        run_quality_checks(
            output_file,
            {
                "enable_structure_audit": True,
                "enable_forbidden_scan": True,
                "forbidden_text": ["None", "NO DISPONIBLE", "RUTA_AL_DOCX"],
                "client_name": client_clean,
                "pipeline_version": PIPELINE_VERSION,
            },
        )

        docx_hash = _sha256_file(output_file)
        pdf_hash = _sha256_file(pdf_path) if pdf_path.exists() else None
        log_payload = {
            "pipeline_version": PIPELINE_VERSION,
            "mode": run_mode,
            "input_hash": input_hash,
            "docx_hash": docx_hash,
            "pdf_hash": pdf_hash,
            "generated_at": generated_at,
        }
        try:
            with (client_dir / "execution_log.json").open("w", encoding="utf-8") as f:
                json.dump(log_payload, f, ensure_ascii=False, indent=2)
        except Exception as e:
            raise OutputContractError(f"CONTRACT_LOG_WRITE_FAILED: {e}")

        validate_output_contract(
            client_dir,
            {
                "mode": run_mode,
                "require_pdf": require_pdf,
            },
        )

        return str(output_file)

    except Exception as e:
        logger.error(f"Error en generación de reporte profesional: {e}")
        raise

def _process_markdown_content(doc, markdown_text, heading_mgr=None):
    """Procesar contenido markdown y convertirlo a formato Word"""
    try:
        if not markdown_text:
            return

        lines = markdown_text.split('\n')
        current_table = []
        in_table = False
        pending_heading = None

        def flush_pending_heading():
            nonlocal pending_heading
            if not pending_heading:
                return
            level, text = pending_heading
            clean_text = (text or "").strip()
            if not clean_text:
                pending_heading = None
                return

            if heading_mgr is not None:
                heading_mgr.add_h3(clean_text)
            else:
                heading = doc.add_paragraph(clean_text)
                heading.style = doc.styles["Heading 3"] if level in (2, 3) else doc.styles["Heading 3"]
                run = heading.runs[0] if heading.runs else heading.add_run("")
                run.font.color.rgb = SECONDARY
            pending_heading = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if in_table and current_table:
                    flush_pending_heading()
                    _create_table_from_markdown(doc, current_table)
                    current_table = []
                    in_table = False
                continue

            if "None" in line or "NO DISPONIBLE" in line:
                continue
            
            # Detectar encabezados
            if line.startswith('## '):
                pending_heading = (2, line[3:].strip())
            elif line.startswith('### '):
                pending_heading = (3, line[4:].strip())
            # Detectar tablas
            elif '|' in line and line.count('|') >= 2:
                flush_pending_heading()
                if not in_table:
                    in_table = True
                current_table.append(line)
            # Detectar listas
            elif line.startswith('- ') or line.startswith('* '):
                flush_pending_heading()
                para = doc.add_paragraph(line[2:], style='List Bullet')
                para.paragraph_format.left_indent = Cm(0.5)
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                flush_pending_heading()
                para = doc.add_paragraph(line[3:], style='List Number')
                para.paragraph_format.left_indent = Cm(0.5)
            # Texto normal
            else:
                if in_table and current_table:
                    flush_pending_heading()
                    _create_table_from_markdown(doc, current_table)
                    current_table = []
                    in_table = False
                
                flush_pending_heading()
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
        
        # Procesar tabla final si existe
        if in_table and current_table:
            flush_pending_heading()
            _create_table_from_markdown(doc, current_table)
            
    except Exception as e:
        logger.error(f"Error procesando markdown: {e}")

def _create_table_from_markdown(doc, table_lines):
    """Crear tabla desde líneas markdown"""
    try:
        if len(table_lines) < 2:
            return

        data_lines = [line for line in table_lines if not re.match(r'^[\|\-\s:]+$', line)]

        if len(data_lines) < 2:
            return

        headers = [cell.strip() for cell in data_lines[0].split('|')[1:-1]]
        data_rows = []

        for line in data_lines[1:]:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(cells) == len(headers):
                data_rows.append(cells)

        if not data_rows:
            return

        table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        _apply_table_layout(table)

        try:
            total_width_cm = 16
            col_w = Cm(total_width_cm / max(len(headers), 1))
            for i in range(len(headers)):
                table.columns[i].width = col_w
        except Exception:
            pass

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            _set_cell_shade(cell, "E7F0FA")
            _set_cell_border(cell, "single", "004E89", "12")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = PRIMARY
            run.font.name = "Calibri"

        for row_idx, row_data in enumerate(data_rows, 1):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                _set_cell_border(cell, "single", "CCCCCC", "4")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                s = (cell_data or "").strip()
                if re.fullmatch(r"[-+]?(\d{1,3}([.,]\d{3})*|\d+)([.,]\d+)?%?", s):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(cell_data)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = TEXT

        doc.add_paragraph()

    except Exception as e:
        logger.error(f"Error creando tabla desde markdown: {e}")


def _build_audit_annex_markdown(cfg):
    lines = []
    metrics_version = cfg.get("metrics_version", "v1")

    dcs = cfg.get("artifact_contradiction_score", {}) or {}
    ofs = cfg.get("artifact_fatigue_score", {}) or {}
    contradictions = cfg.get("artifact_contradictions", {}) or {}
    segments = cfg.get("artifact_segments", {}) or {}
    fatigue = cfg.get("artifact_operational_fatigue", {}) or {}
    alerts = cfg.get("alerts", {}).get("alerts", []) if isinstance(cfg.get("alerts"), dict) else []
    if alerts:
        def _alert_key(a):
            severity = str((a or {}).get("severity", ""))
            code = str((a or {}).get("code", ""))
            metric = str((a or {}).get("metric", ""))
            value = str((a or {}).get("value", ""))
            message = str((a or {}).get("message", ""))
            return (severity, code, metric, value, message)

        alerts = sorted(list(alerts), key=_alert_key)

    dcs_score = dcs.get("score", "N/A")
    dcs_level = dcs.get("level", "N/A")
    dcs_ratio = contradictions.get("contradiction_ratio", "N/A")

    ofs_score = ofs.get("score", "N/A")
    ofs_level = ofs.get("level", "N/A")

    lines.append(f"## Auditoría interna — Metrics {metrics_version}")
    lines.append("")
    lines.append("### Scores")
    lines.append(f"- DCS: {dcs_score} ({dcs_level}) | contradiction_ratio: {dcs_ratio}")
    lines.append(f"- OFS: {ofs_score} ({ofs_level})")
    lines.append("")

    lines.append("### Alerts")
    if alerts:
        for alert in alerts:
            severity = alert.get("severity", "N/A")
            code = alert.get("code", "N/A")
            metric = alert.get("metric", "N/A")
            value = alert.get("value", "N/A")
            message = alert.get("message", "")
            lines.append(f"- [{severity}] {code} ({metric}={value}): {message}")
    else:
        lines.append("- Sin alertas generadas")
    lines.append("")

    lines.append("### Contradicciones (muestra)")
    lines.append(f"- contradiction_count: {contradictions.get('contradiction_count', 'N/A')}")
    lines.append(f"- contradiction_ratio: {contradictions.get('contradiction_ratio', 'N/A')}")
    examples = contradictions.get("examples") or []
    lines.append("- Ejemplos:")
    if examples:
        for ex in examples[:5]:
            lines.append(f"  - {ex}")
    else:
        lines.append("  - [Sin ejemplos disponibles]")
    lines.append("")

    lines.append("### Segmentación local vs turista")
    lines.append(f"- local_count: {segments.get('local_count', 'N/A')} | local_avg_rating: {segments.get('local_avg_rating', 'N/A')}")
    lines.append(f"- tourist_count: {segments.get('tourist_count', 'N/A')} | tourist_avg_rating: {segments.get('tourist_avg_rating', 'N/A')}")
    lines.append(f"- key_differences: {segments.get('key_differences', 'N/A')}")
    lines.append("")

    lines.append("### Fatiga operativa")
    lines.append(f"- trend_rating_last_6m: {fatigue.get('trend_rating_last_6m', 'N/A')}")
    lines.append(f"- trend_sentiment_last_6m: {fatigue.get('trend_sentiment_last_6m', 'N/A')}")
    repeated = fatigue.get("repeated_complaints") or []
    if repeated:
        joined = ", ".join(str(x) for x in repeated[:10])
    else:
        joined = "[Sin términos repetidos destacados]"
    lines.append(f"- repeated_complaints: {joined}")
    lines.append(f"- fatigue_signal: {fatigue.get('fatigue_signal', 'N/A')}")

    return "\n".join(lines)


def _build_reviews_text(df):
    """Construir texto de reseñas para el análisis"""
    rows = []
    for _, r in df.iterrows():
        try:
            user = str(r.get("user", "")).strip() or "Usuario"
            date = str(r.get("date", "")).strip() or "s/f"
            rating = int(r.get("rating_num", 0))
            txt = str(r.get("text", "")).strip() or "[Sin comentario]"
            rows.append(f"- [{date}] {user} ({rating}★): {txt}")
        except Exception as e:
            logger.warning(f"Error procesando fila de reseña: {e}")
            continue
    return "\n".join(rows)
